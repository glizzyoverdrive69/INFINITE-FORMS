#!/usr/bin/env python3
"""
Infinite Forms Plugin
======================
A floating control panel for DaVinci Resolve Studio, built with Fusion's
UIManager. Runs as its own window alongside Resolve rather than docking
into it -- there's no true panel-docking system for third-party tools,
so this behaves like a tool palette that floats on top.

Install (macOS):
  1. Copy this file into:
     ~/Library/Application Support/Blackmagic Design/DaVinci Resolve/Fusion/Scripts/Utility/
     (create the "Utility" folder if it doesn't already exist)
  2. In Resolve: Preferences > General > External scripting using -> Local
  3. Workspace > Scripts > Utility > infinite_forms_plugin
     (Resolve lists scripts by filename, not by the window title below)

Requires DaVinci Resolve Studio -- the UIManager used here isn't available
in the free version.
"""

BUILD_TAG = "2026-08-05.1"
print(f"[Infinite Forms] script starting -- build {BUILD_TAG}")

# --- Auto-update -------------------------------------------------------
# The GitHub "username/repo" the update check reads from. The repo must be
# PUBLIC and contain a VERSION file (whose content matches the plugin's
# BUILD_TAG) plus the plugin file itself, both at the repo root on
# UPDATE_BRANCH. Leave UPDATE_REPO empty to disable all update behaviour,
# including the panel's Check for Update button.
UPDATE_REPO = "glizzyoverdrive69/INFINITE-FORMS"
UPDATE_BRANCH = "main"
UPDATE_PLUGIN_FILENAME = "infinite_forms_plugin.py"

# Every launch checks for an update automatically. When that check finds
# something worth acting on -- a newer build, a private repo, missing
# certificates -- it pops the same dialog the Check for Update button
# shows, so it cannot be missed. Set to False to keep launch findings in
# the panel log only. Being up to date, or simply offline, never pops
# anything either way.
UPDATE_STARTUP_DIALOG = True

import importlib.machinery
import importlib.util
import os
from difflib import SequenceMatcher
import re
import ssl
import sys
import time
import traceback
import unicodedata
import urllib.error
import urllib.request

# ---------------------------------------------------------------------------
# Connect to Resolve / Fusion.
#
# Layered bootstrap, because "import DaVinciResolveScript" isn't
# available on every install (this machine's original scripts loaded
# fusionscript.so directly for exactly this reason):
#   1. use the `bmd` object Resolve injects into menu scripts, if any
#   2. plain import of DaVinciResolveScript
#   3. add the standard Modules folders to sys.path, retry the import
#   4. last resort: load fusionscript.so directly (RESOLVE_SCRIPT_LIB
#      or the default app path)
# ---------------------------------------------------------------------------

def _load_bmd():
    try:
        import DaVinciResolveScript as mod
        return mod
    except ImportError:
        pass

    candidates = [
        os.path.join(os.environ.get("RESOLVE_SCRIPT_API", ""), "Modules"),
        "/Library/Application Support/Blackmagic Design/DaVinci Resolve/Developer/Scripting/Modules",
        os.path.expanduser(
            "~/Library/Application Support/Blackmagic Design/DaVinci Resolve/Developer/Scripting/Modules"
        ),
    ]
    for path in candidates:
        if path and os.path.isdir(path) and path not in sys.path:
            sys.path.append(path)
    try:
        import DaVinciResolveScript as mod
        return mod
    except ImportError:
        pass

    lib_path = os.environ.get("RESOLVE_SCRIPT_LIB") or (
        "/Applications/DaVinci Resolve/DaVinci Resolve.app"
        "/Contents/Libraries/Fusion/fusionscript.so"
    )
    if os.path.isfile(lib_path):
        loader = importlib.machinery.ExtensionFileLoader("fusionscript", lib_path)
        spec = importlib.util.spec_from_loader("fusionscript", loader)
        mod = importlib.util.module_from_spec(spec)
        loader.exec_module(mod)
        if hasattr(mod, "scriptapp"):
            return mod

    raise RuntimeError(
        "Could not locate Resolve's scripting module -- checked the"
        " DaVinciResolveScript import, the standard Modules folders, and"
        " fusionscript.so inside the Resolve app."
    )


try:
    bmd  # noqa: B018 -- injected by Resolve when run from the Scripts menu
except NameError:
    bmd = _load_bmd()

try:
    resolve  # noqa: B018 -- also injected on most installs
except NameError:
    resolve = bmd.scriptapp("Resolve")

fusion = resolve.Fusion()
project_manager = resolve.GetProjectManager()
print("[Infinite Forms] connected to Resolve")


try:
    PLUGIN_FILE_PATH = os.path.abspath(__file__)
except Exception:
    PLUGIN_FILE_PATH = os.path.expanduser(
        "~/Library/Application Support/Blackmagic Design/DaVinci Resolve"
        "/Fusion/Scripts/Utility/infinite_forms_plugin.py")


def _ensure_user_site_packages():
    """pip --user installs (python-docx) land in the user site-packages
    folder, but Resolve's embedded Python doesn't include that folder on
    its import path on every machine. Add it when it exists -- harmless
    where it's already present."""
    candidates = []
    try:
        import site
        candidates.append(site.getusersitepackages())
    except Exception:
        pass
    version_tag = f"{sys.version_info.major}.{sys.version_info.minor}"
    candidates.append(os.path.expanduser(
        f"~/Library/Python/{version_tag}/lib/python/site-packages"))
    for path in candidates:
        if path and os.path.isdir(path) and path not in sys.path:
            sys.path.insert(0, path)


_ensure_user_site_packages()

ui = fusion.UIManager
disp = bmd.UIDispatcher(ui)


def get_context():
    """Re-fetch project/timeline/media pool on every click, so the panel
    always acts on whatever is currently open in Resolve -- not whatever
    was open when the panel first launched."""
    project = project_manager.GetCurrentProject()
    if not project:
        return None, None, None
    timeline = project.GetCurrentTimeline()
    media_pool = project.GetMediaPool()
    return project, timeline, media_pool


# ---------------------------------------------------------------------------
# Config -- camera audio replace (from Replace_Camera_Audio.py)
# ---------------------------------------------------------------------------
AUDIO_VIDEO_TRACK_INDEX = 1
AUDIO_DEST_TRACK_NAME = "CAM RESTORE (mono)"
AUDIO_DEST_TRACK_TYPE = "mono"
AUDIO_ONLY_IF_NO_LINKED_AUDIO = True
AUDIO_LINK_RESTORED_AUDIO = True
AUDIO_DRY_RUN = False

# ---------------------------------------------------------------------------
# Config -- Writer's Reel location lower thirds (from
# create_neighbourhood_lower_thirds.py)
# ---------------------------------------------------------------------------
LT_SOURCE_VIDEO_TRACKS = [1, 2, 3, 4]   # read clips from V1-V4
LT_TITLE_VIDEO_TRACK = 5                 # stamp titles on V5, above them all
LT_TITLE_TEMPLATE_NAME = "TM_LOWER_LEFT_THIRD_TEMPLATE"
LT_TEXT_TOOL_NAME = "Template"
LT_GENERATED_TITLE_PREFIX = "AUTO_NEIGHBOURHOOD_LT__"
LT_UPDATE_EXISTING_GENERATED_TITLES = True
LT_DRY_RUN = False
LT_DEBUG_TEXT_TOOLS = False

# ---------------------------------------------------------------------------
# New-timeline settings -- applied to every timeline the assemble pipeline
# creates. Format and color fields are confirmed against Resolve's
# scripting API; a few obscure Monitor/Output fields aren't reliably
# documented, so they're left out until verified against a real
# GetSetting() dump rather than guessed at.
# ---------------------------------------------------------------------------
TIMELINE_SETTINGS = {
    # Must be set FIRST: new timelines default to "Use Project Settings",
    # and resolution/framerate keys silently fail until it's disabled.
    "useCustomSettings": "1",

    "timelineResolutionWidth": "3840",
    "timelineResolutionHeight": "2160",
    "timelineFrameRate": "29.97",
    "timelinePixelAspectRatio": "square",  # confirmed via real run log

    "colorScienceMode": "davinciYRGB",
    "colorSpaceTimeline": "DaVinci WG/Intermediate",
    "colorSpaceOutput": "Rec.709-A",
}

# Gap inside each assembly group, separating the approved Clip Asset
# Package clips from the leftover bin clips that follow them.
ASSEMBLY_EXTRAS_GAP_SECONDS = 15

# ---------------------------------------------------------------------------
# Config -- client colour groups. The .drx files are grades exported from
# a Gallery still (right-click a still > Export -- the .drx lands next to
# the image file). Fill in the real paths; empty string skips the
# clip-grade step and only does group assignment + settings check.
#
# Pre/post-clip GROUP grades can't be authored via the API -- they're
# per-project, set up once by a colorist (ideally in the template project
# your jobs start from). Assigning clips to the group inherits them
# automatically. If a group is missing here, it gets created empty and
# the log says the pre/post grades still need their one-time setup.
# ---------------------------------------------------------------------------
CLIENT_COLOR_PRESETS = {
    "EXPEDIA": {
        "group_name": "CRM - Expedia",
        "drx_path": "",  # no clip-grade still exists yet -- step is skipped
    },
    "SKYSCANNER": {
        "group_name": "CRM - Skyscanner",
        "drx_path": "",  # no clip-grade still exists yet -- step is skipped
    },
}

# ---------------------------------------------------------------------------
# Config -- standard track layout. Missing tracks are created so the
# full layout always exists; audio tracks are created as the subtype
# given per track ("mono"/"stereo").
# ---------------------------------------------------------------------------
VIDEO_TRACK_NAMES = {
    1: "VIDEO A",
    2: "VIDEO B",
    3: "GRAPHICS",
}

AUDIO_TRACK_NAMES = {
    1: ("CAM AUDIO", "mono"),
    2: ("MUSIC 1", "stereo"),
    3: ("MUSIC 2", "stereo"),
    4: ("TEMP VO", "mono"),
    5: ("HUMAN VO", "mono"),
    6: ("SFX 1", "stereo"),
    7: ("SFX 2", "stereo"),
    8: ("MASTER", "stereo"),
}


# ---------------------------------------------------------------------------
# Panel theme -- QSS matching the Infinite Forms mockup (forest green +
# brass/gold). UIManager passes unknown properties through to Qt on most
# builds; where StyleSheet isn't honoured, the panel simply stays plain
# and everything still works.
# ---------------------------------------------------------------------------
PANEL_QSS = """
QWidget { background-color: #15221A; color: #E7E3D3; font-size: 13px; }
QPushButton {
    background-color: #1C2B21; color: #C4A253;
    border: 1px solid #8C7238; border-radius: 6px; padding: 8px 12px;
}
QPushButton:hover { border-color: #C4A253; background-color: #223327; }
QComboBox, QLineEdit, QTextEdit {
    background-color: #0E1710; color: #E7E3D3;
    border: 1px solid #2B3E30; border-radius: 4px; padding: 4px 6px;
}
QLabel { background: transparent; }
"""

# Optional custom texture: set an image path and it tiles behind the
# panel and every dialog (depends on this build honouring stylesheets
# at all -- same caveat as the colour theme).
PANEL_TEXTURE_PATH = ""  # e.g. "/Users/you/Pictures/panel_texture.png"
if PANEL_TEXTURE_PATH and os.path.isfile(PANEL_TEXTURE_PATH):
    PANEL_QSS += (
        '\nQWidget#AutomationPanel, QWidget {'
        f' background-image: url("{PANEL_TEXTURE_PATH}");'
        ' background-repeat: repeat; }'
    )

HEADER_HTML = (
    '<div align="center">'
    '<span style="color:#C4A253; font-size:24px; font-family:Georgia,serif;'
    ' letter-spacing:6px;"><b>INFINITE FORMS</b></span><br>'
    '<span style="color:#8A9384; font-size:10px; letter-spacing:3px;">'
    'TOOLS TO WORK SMARTER, NOT HARDER</span></div>'
)


def section_label(text):
    return ui.Label({
        "Text": f'<span style="color:#C4A253; font-size:11px;'
                f' letter-spacing:2px;"><b>{text.upper()}</b></span>'
    })


def apply_panel_style(window):
    """Best-effort theming -- harmless no-op on builds that ignore it."""
    try:
        window.StyleSheet = PANEL_QSS
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Window layout -- rebuildable, so Pin (always-on-top) can recreate it.
# Buttons are listed alphabetically by feature name.
# ---------------------------------------------------------------------------
PANEL = {"collapsed": False, "pinned": False}
win = None
items = None

PANEL_FULL_SIZE = [480, 700]
PANEL_COLLAPSED_SIZE = [480, 84]


def build_main_panel():
    """(Re)create the main panel window. Globals win/items always point
    at the current instance, so log() and handlers keep working across
    rebuilds (used by the Pin toggle)."""
    global win, items

    props = {
        "ID": "AutomationPanel",
        "WindowTitle": "Infinite Forms",
        "Geometry": [100, 100] + (PANEL_COLLAPSED_SIZE if PANEL["collapsed"]
                                  else PANEL_FULL_SIZE),
        "StyleSheet": PANEL_QSS,
    }
    if PANEL["pinned"]:
        props["WindowFlags"] = {"Window": True, "WindowStaysOnTopHint": True}

    new_win = disp.AddWindow(
        props,
        [
            ui.VGroup(
                {"Spacing": 10},
                [
                    ui.HGroup({"Spacing": 6, "Weight": 0}, [
                        ui.Label({"Text": HEADER_HTML, "Weight": 1}),
                    ] + ([
                        ui.Button({"ID": "BtnUpdate",
                                   "Text": f"Update ({UPDATE_STATE['remote']})",
                                   "Weight": 0}),
                    ] if UPDATE_STATE["available"] else []) + [
                        ui.Button({"ID": "BtnPin",
                                   "Text": "Unpin" if PANEL["pinned"] else "Pin",
                                   "Weight": 0}),
                        ui.Button({"ID": "BtnCollapse",
                                   "Text": "+" if PANEL["collapsed"] else "\u2013",
                                   "Weight": 0}),
                    ]),
                    ui.VGroup(
                        {"ID": "PanelBody", "Spacing": 10},
                        [
                            ui.Button({"ID": "BtnLowerThirds", "Text": "Auto Lower Thirds"}),
                            ui.Button({"ID": "BtnBinFinder", "Text": "Bin Finder"}),
                            ui.Button({"ID": "BtnApplyClientColor", "Text": "Colour Grading Prep"}),
                            ui.Button({"ID": "BtnAssemble", "Text": "Mid/Short Form Assembly"}),
                            ui.Button({"ID": "BtnRenameTracks", "Text": "Rename Video & Audio Tracks"}),
                            ui.Button({"ID": "BtnAudioSync", "Text": "Replace Camera Audio"}),
                            ui.Button({"ID": "BtnSortShootNotes", "Text": "Sort by Shoot Notes"}),
                            ui.TextEdit({"ID": "Log", "ReadOnly": True,
                                         "Text": "Ready.\n", "Weight": 1}),
                            # Footer: which build this is, and a manual
                            # update check. Sits below the log so it never
                            # competes with the feature buttons for
                            # attention.
                            ui.HGroup({"Spacing": 6, "Weight": 0}, [
                                ui.Label({
                                    "ID": "BuildLabel",
                                    "Text": f'<span style="color:#8A9384;'
                                            f' font-size:10px;'
                                            f' letter-spacing:1px;">BUILD'
                                            f' {BUILD_TAG}</span>',
                                    "Weight": 1,
                                }),
                                ui.Button({"ID": "BtnCheckUpdate",
                                           "Text": "Check for Update",
                                           "Weight": 0}),
                            ]),
                        ],
                    ),
                ],
            )
        ],
    )

    win = new_win
    items = win.GetItems()
    _flush_log_widget()

    # Preemptive no-op handlers for chatty widget events, so nothing the
    # Log box (or anything else) emits can reach the dispatcher
    # unregistered and trigger KeyError: 'On'.
    def _noop(_ev=None):
        pass
    for event_name in ("TextChanged",):
        try:
            setattr(win.On.Log, event_name, _noop)
        except Exception:
            pass

    win.On.AutomationPanel.Close = on_close
    if UPDATE_STATE["available"]:
        try:
            win.On.BtnUpdate.Clicked = guard(on_apply_update)
        except Exception:
            pass
    win.On.BtnPin.Clicked = guard(on_toggle_pin)
    win.On.BtnCollapse.Clicked = guard(on_toggle_collapse)
    win.On.BtnCheckUpdate.Clicked = guard(on_check_for_update)
    win.On.BtnAudioSync.Clicked = guard(on_audio_sync)
    win.On.BtnLowerThirds.Clicked = guard(on_lower_thirds)
    win.On.BtnAssemble.Clicked = guard(on_assemble)
    win.On.BtnApplyClientColor.Clicked = guard(on_apply_client_color)
    win.On.BtnRenameTracks.Clicked = guard(on_rename_tracks)
    win.On.BtnSortShootNotes.Clicked = guard(on_sort_shoot_notes)
    win.On.BtnBinFinder.Clicked = guard(on_bin_finder)

    apply_panel_style(win)
    _apply_collapse_state()
    win.Show()
    print("[Infinite Forms] panel constructed")
    return win


def _apply_collapse_state():
    try:
        items["PanelBody"].Hidden = PANEL["collapsed"]
    except Exception:
        pass
    try:
        win.Resize(PANEL_COLLAPSED_SIZE if PANEL["collapsed"] else PANEL_FULL_SIZE)
        win.RecalcLayout()
    except Exception:
        pass
    try:
        items["BtnCollapse"].Text = "+" if PANEL["collapsed"] else "\u2013"
    except Exception:
        pass


def on_toggle_collapse(_ev=None):
    """Resize() is a no-op on this build, so collapsing rebuilds the
    window with the small geometry baked in -- the same (proven)
    mechanism Pin uses."""
    PANEL["collapsed"] = not PANEL["collapsed"]
    old_win = win
    try:
        old_win.Hide()
    except Exception:
        pass
    build_main_panel()


def on_toggle_pin(_ev=None):
    """Always-on-top needs window flags, which only apply at creation on
    this toolkit -- so pinning rebuilds the panel with the flag set."""
    PANEL["pinned"] = not PANEL["pinned"]
    old_win = win
    try:
        old_win.Hide()
    except Exception:
        pass
    build_main_panel()
    log("Panel pinned on top." if PANEL["pinned"] else "Panel unpinned.")


LOG_LINES = ["Ready."]
_LOG_WIDGET_HOLD = {"depth": 0}


def _flush_log_widget():
    try:
        items["Log"].Text = "\n".join(LOG_LINES) + "\n"
    except Exception:
        pass  # console print still has the full trail


def hold_log_widget():
    """Writing to the log widget emits an event owned by the MAIN
    dispatcher; if a nested dialog's dispatcher is running, that event
    crashes it (KeyError: 'On') -- and logging the crash writes to the
    log again, feeding an infinite error loop. So while any blocking
    dialog is open, log() buffers only; the widget updates on release."""
    _LOG_WIDGET_HOLD["depth"] += 1


def release_log_widget():
    _LOG_WIDGET_HOLD["depth"] = max(0, _LOG_WIDGET_HOLD["depth"] - 1)
    if _LOG_WIDGET_HOLD["depth"] == 0:
        _flush_log_widget()


def log(message):
    """Append a line to the on-panel log box, and print it too, so
    Workspace > Console shows the same trail. The log is kept in Python
    and only ever written to the widget -- reading .Text back returns
    None on some Resolve builds, which breaks string appends."""
    LOG_LINES.append(str(message))
    print(message)
    if _LOG_WIDGET_HOLD["depth"] == 0:
        _flush_log_widget()


def guard(handler):
    """Wrap a UI event handler so an exception logs a traceback instead
    of unwinding the event loop -- an uncaught error in any handler
    otherwise kills the whole panel (every window vanishes)."""
    def wrapped(ev=None):
        try:
            return handler(ev)
        except Exception:
            log(f"Error in {getattr(handler, '__name__', 'handler')}:\n"
                f"{traceback.format_exc()}")
    return wrapped


def log_quiet(message):
    """Log to the Python buffer + Console only -- never the widget.
    Used inside dispatcher-error handling, where a widget write could
    emit the very kind of event that caused the error."""
    LOG_LINES.append(str(message))
    print(message)


def run_loop_resilient(dispatcher, context=""):
    """Run a dispatcher loop, surviving errors thrown inside Resolve's
    OWN event dispatch (e.g. KeyError: 'On' when a widget emits an event
    nothing registered for). Such errors otherwise collapse every window.
    Logs and resumes; gives up after repeated failures."""
    failures = 0
    while True:
        try:
            dispatcher.RunLoop()
            return
        except Exception:
            failures += 1
            log_quiet(f"Dispatcher error{f' in {context}' if context else ''}"
                      f" (resumed):\n{traceback.format_exc()}")
            if failures >= 25:
                log_quiet("Too many dispatcher errors -- closing this loop.")
                return


# ---------------------------------------------------------------------------
# Camera audio replace -- helpers
# ---------------------------------------------------------------------------
def ensure_audio_track(target_timeline):
    audio_count = target_timeline.GetTrackCount("audio")
    for i in range(1, audio_count + 1):
        try:
            if target_timeline.GetTrackName("audio", i) == AUDIO_DEST_TRACK_NAME:
                return i
        except Exception:
            pass

    ok = target_timeline.AddTrack("audio", AUDIO_DEST_TRACK_TYPE)
    if not ok:
        raise RuntimeError("Could not create mono audio track.")

    new_idx = target_timeline.GetTrackCount("audio")
    try:
        target_timeline.SetTrackName("audio", new_idx, AUDIO_DEST_TRACK_NAME)
    except Exception:
        pass
    return new_idx


def has_linked_audio(video_item):
    try:
        linked = video_item.GetLinkedItems() or []
    except Exception:
        return False

    for item in linked:
        try:
            track_type, _track_index = item.GetTrackTypeAndIndex()
            if track_type == "audio":
                return True
        except Exception:
            continue
    return False


def get_clip_path(media_pool_item):
    try:
        props = media_pool_item.GetClipProperty() or {}
    except Exception:
        return ""

    for key in ["File Path", "FilePath", "File Name", "Filename", "Clip Name"]:
        value = props.get(key)
        if value:
            return value
    return ""


# ---------------------------------------------------------------------------
# Button 1 -- camera audio replace (ported from Replace_Camera_Audio.py)
# ---------------------------------------------------------------------------
def on_audio_sync(ev):
    """Processes V1 only: restores each clip's own embedded audio onto a
    dedicated mono track and links it back to the video, skipping
    anything that already has linked audio."""
    project, timeline, media_pool = get_context()
    if not timeline:
        log("No timeline open -- open a project and timeline first.")
        return

    try:
        resolve.OpenPage("edit")
    except Exception:
        pass

    video_track_count = timeline.GetTrackCount("video")
    if AUDIO_VIDEO_TRACK_INDEX > video_track_count:
        log(f"Timeline only has {video_track_count} video track(s) -- cannot process V{AUDIO_VIDEO_TRACK_INDEX}.")
        return

    dest_audio_track = ensure_audio_track(timeline)

    try:
        if timeline.GetIsTrackLocked("audio", dest_audio_track):
            log(f"Destination audio track A{dest_audio_track} is locked -- unlock it and try again.")
            return
    except Exception:
        pass

    v_items = timeline.GetItemListInTrack("video", AUDIO_VIDEO_TRACK_INDEX) or []
    log(f"Found {len(v_items)} clip(s) on V{AUDIO_VIDEO_TRACK_INDEX}.")

    restored = skipped = failed = 0

    for video_item in v_items:
        try:
            name = video_item.GetName()

            if AUDIO_ONLY_IF_NO_LINKED_AUDIO and has_linked_audio(video_item):
                skipped += 1
                log(f"Skip (already has linked audio): {name}")
                continue

            media_pool_item = video_item.GetMediaPoolItem()
            if not media_pool_item:
                skipped += 1
                log(f"Skip (no source clip): {name}")
                continue

            clip_path = get_clip_path(media_pool_item)
            record_frame = int(video_item.GetStart(False))
            source_start = int(video_item.GetSourceStartFrame())
            # endFrame is exclusive: derive it from the item's actual
            # duration so restored audio is exactly the video's length.
            # GetSourceEndFrame() here clipped the final frame.
            source_end = source_start + int(video_item.GetDuration(False))

            clip_info = {
                "mediaPoolItem": media_pool_item,
                "startFrame": source_start,
                "endFrame": source_end,
                "mediaType": 2,  # audio only
                "trackIndex": int(dest_audio_track),
                "recordFrame": record_frame,
            }

            if AUDIO_DRY_RUN:
                log(f"DRY RUN -- would restore: {name} ({clip_path})")
                restored += 1
                continue

            new_items = media_pool.AppendToTimeline([clip_info]) or []
            if not new_items:
                failed += 1
                log(f"Failed to append audio for: {name} ({clip_path})")
                continue

            if AUDIO_LINK_RESTORED_AUDIO:
                try:
                    timeline.SetClipsLinked([video_item] + new_items, True)
                except Exception:
                    pass

            restored += 1
            log(f"Restored: {name} ({clip_path})")

        except Exception:
            failed += 1
            log(f"Error on a clip:\n{traceback.format_exc()}")

    log(f"Done. Restored {restored}, skipped {skipped}, failed {failed}.")
    log(f"Audio landed on A{dest_audio_track} ({AUDIO_DEST_TRACK_NAME}).")


# ---------------------------------------------------------------------------
# Writer's Reel location lower thirds -- helpers (ported from
# create_neighbourhood_lower_thirds.py). walk_media_pool_folder() and
# media_pool_item_id() are general enough that the assemble pipeline's
# location-matching step can reuse them too, rather than re-walking the
# bin tree a second way.
# ---------------------------------------------------------------------------
def media_pool_item_id(media_pool_item):
    """A stable identity for a clip -- separate lookups of "the same"
    MediaPoolItem aren't guaranteed to be == to each other."""
    if media_pool_item is None:
        return None
    for method_name in ("GetUniqueId", "GetMediaId"):
        try:
            value = getattr(media_pool_item, method_name)()
            if value:
                return str(value)
        except Exception:
            pass
    return str(id(media_pool_item))


def timeline_item_start(item):
    return int(item.GetStart(False))


def timeline_item_duration(item):
    return int(item.GetDuration(False))


def ensure_video_track(target_timeline, track_index):
    while target_timeline.GetTrackCount("video") < track_index:
        ok = target_timeline.AddTrack("video")
        if not ok:
            raise RuntimeError(f"Could not add video track V{track_index}")


def clean_label(name):
    if not name:
        return ""
    label = name.strip()
    if label.endswith("(N)"):
        label = label[:-3].strip()
    return label


def walk_media_pool_folder(folder, path, clip_folder_map, all_clips_by_name,
                            clip_by_id=None):
    """Recursively index the whole Media Pool bin tree once: clip id ->
    full folder path, clip name -> clip (used to find the title
    template), and optionally clip id -> clip object (used by the
    assembly's fill-with-unused-bin-clips step)."""
    folder_name = folder.GetName()
    current_path = path + [folder_name]

    for clip in folder.GetClipList():
        clip_id = media_pool_item_id(clip)
        if clip_id:
            clip_folder_map[clip_id] = current_path
            if clip_by_id is not None:
                clip_by_id[clip_id] = clip
        try:
            clip_name = clip.GetName()
            all_clips_by_name.setdefault(clip_name, []).append(clip)
        except Exception:
            pass

    for subfolder in folder.GetSubFolderList():
        walk_media_pool_folder(subfolder, current_path, clip_folder_map,
                                all_clips_by_name, clip_by_id)


def find_template_clip_by_name(all_clips_by_name, template_name):
    matches = all_clips_by_name.get(template_name, [])
    if not matches:
        return None
    if len(matches) > 1:
        log(f"Warning: found {len(matches)} Media Pool items named '{template_name}'. Using the first one.")
    return matches[0]


def label_from_folder_path(folder_path):
    """Find the nearest folder ending in (N) -- a neighbourhood/region
    marker -- then use the folder directly below it as the label. E.g.
    Master / Northwest (N) / Dupont Circle / TM - Photographer Name
    -> "Dupont Circle". Falls back to the clip's parent bin if no (N)
    marker exists anywhere in the path."""
    if not folder_path:
        return ""

    last_marker_index = None
    for index, folder_name in enumerate(folder_path):
        if folder_name.strip().endswith("(N)"):
            last_marker_index = index

    if last_marker_index is not None:
        candidate_index = last_marker_index + 1
        if candidate_index < len(folder_path):
            return clean_label(folder_path[candidate_index])
        return clean_label(folder_path[last_marker_index])

    if len(folder_path) >= 2:
        return clean_label(folder_path[-2])
    return clean_label(folder_path[-1])


def get_tool_name(tool):
    try:
        attrs = tool.GetAttrs()
        return attrs.get("TOOLS_Name", "")
    except Exception:
        return ""


def get_text_tools_from_comp(comp):
    found = []
    try:
        tools = comp.GetToolList(False)
    except Exception:
        return found

    for _, tool in tools.items():
        tool_name = get_tool_name(tool)
        if tool_name.startswith("MediaOut"):
            continue
        try:
            current_text = tool.GetInput("StyledText")
            found.append({"name": tool_name, "tool": tool, "current_text": current_text})
        except Exception:
            pass

    found.sort(key=lambda item: item["name"])
    return found


def print_detected_text_tools(title_name, text_tools):
    print(f"\nDetected StyledText tools for title '{title_name}':")
    if not text_tools:
        print("  No Text+ / StyledText tools found.")
    else:
        for index, item in enumerate(text_tools, start=1):
            print(f"  {index}. name='{item['name']}' current_text={item['current_text']!r}")
    print("")


def find_text_tool(text_tools):
    for item in text_tools:
        if item["name"] == LT_TEXT_TOOL_NAME:
            return item["tool"]

    if len(text_tools) == 1:
        print(f"Warning: expected Text+ node '{LT_TEXT_TOOL_NAME}', using the only detected node '{text_tools[0]['name']}'.")
        return text_tools[0]["tool"]

    print(f"\nCould not find the required Text+ node '{LT_TEXT_TOOL_NAME}'. Detected nodes:")
    for item in text_tools:
        print(f"  - {item['name']}")
    print("")
    return None


def set_textplus_text(timeline_item, label_text):
    """Set a title's Text+ field. Resolve can take a moment to expose the
    Fusion comp right after appending a title, so this retries for up to
    ~6 seconds rather than failing immediately."""
    title_name = timeline_item.GetName()
    comp = None

    for _ in range(30):
        try:
            comp_count = timeline_item.GetFusionCompCount()
            if comp_count and comp_count > 0:
                comp = timeline_item.GetFusionCompByIndex(1)
                if comp:
                    break
        except Exception:
            pass
        time.sleep(0.2)

    if not comp:
        print(f"Could not access Fusion comp for title '{title_name}'.")
        print("This usually means the template isn't a Fusion/Text+ title, or Resolve hasn't exposed the comp yet.")
        return False

    text_tools = get_text_tools_from_comp(comp)
    if LT_DEBUG_TEXT_TOOLS:
        print_detected_text_tools(title_name, text_tools)

    text_tool = find_text_tool(text_tools)
    if not text_tool:
        print_detected_text_tools(title_name, text_tools)
        return False

    try:
        text_tool.SetInput("StyledText", label_text)
        return True
    except Exception:
        print(f"Failed while setting Text+ field for '{title_name}'.")
        traceback.print_exc()
        return False


def existing_generated_titles_by_start(target_timeline, title_track):
    items_by_start = {}
    if title_track > target_timeline.GetTrackCount("video"):
        return items_by_start
    for item in target_timeline.GetItemListInTrack("video", title_track):
        try:
            name = item.GetName() or ""
            if name.startswith(LT_GENERATED_TITLE_PREFIX):
                items_by_start[timeline_item_start(item)] = item
        except Exception:
            pass
    return items_by_start


def apply_lower_third(media_pool, existing_titles, template_clip, label_text,
                       start_frame, duration, title_track):
    """Create or update ONE lower third at a timeline position -- the
    atomic operation. Both the standalone button below and the future
    assemble pipeline call this per clip/group once they've resolved a
    label and a position, rather than duplicating the Text+ mechanics."""
    title_name = LT_GENERATED_TITLE_PREFIX + label_text
    existing_title_item = existing_titles.get(start_frame)

    if existing_title_item and LT_UPDATE_EXISTING_GENERATED_TITLES:
        try:
            existing_title_item.SetName(title_name)
        except Exception:
            pass
        return set_textplus_text(existing_title_item, label_text)

    clip_info = {
        "mediaPoolItem": template_clip,
        "startFrame": 0,
        "endFrame": duration,  # frame-length fix: not duration - 1
        "trackIndex": title_track,
        "recordFrame": start_frame,
        "mediaType": 1,
    }
    new_items = media_pool.AppendToTimeline([clip_info])
    if not new_items:
        return False

    title_item = new_items[0]
    try:
        title_item.SetName(title_name)
    except Exception:
        pass

    ok = set_textplus_text(title_item, label_text)
    if ok:
        existing_titles[start_frame] = title_item
    return ok


# ---------------------------------------------------------------------------
# Confirmation dialog -- blocks until the user picks Run or Cancel, using
# its own dispatcher so the main panel loop is undisturbed.
# ---------------------------------------------------------------------------
def confirm_dialog(title, message):
    dlg_disp = bmd.UIDispatcher(ui)
    result = {"ok": False}

    dlg = dlg_disp.AddWindow(
        {
            "ID": "ConfirmDlg",
            "WindowTitle": title,
            "Geometry": [200, 200, 420, 150],
            "StyleSheet": PANEL_QSS,
        },
        [
            ui.VGroup(
                {"Spacing": 10},
                [
                    ui.Label({"Text": message, "WordWrap": True, "Weight": 1}),
                    ui.HGroup(
                        {"Spacing": 8, "Weight": 0},
                        [
                            ui.Button({"ID": "BtnDlgCancel", "Text": "Cancel"}),
                            ui.Button({"ID": "BtnDlgRun", "Text": "Run"}),
                        ],
                    ),
                ],
            )
        ],
    )

    def on_run(_ev):
        result["ok"] = True
        dlg_disp.ExitLoop()

    def on_cancel(_ev):
        dlg_disp.ExitLoop()

    dlg.On.BtnDlgRun.Clicked = on_run
    dlg.On.BtnDlgCancel.Clicked = on_cancel
    dlg.On.ConfirmDlg.Close = on_cancel

    hold_log_widget()
    dlg.Show()
    run_loop_resilient(dlg_disp, "confirm dialog")
    dlg.Hide()
    release_log_widget()
    return result["ok"]


# ---------------------------------------------------------------------------
# Lower-thirds job runner. Timer-driven: one clip per tick, so the event
# loop breathes between clips -- the progress window repaints and the
# Stop button actually receives its click. (A plain loop can't do this:
# StepLoop, the usual event pump, errors when called from Python.) If
# the Timer API isn't available on this build, falls back to a
# synchronous run and says so honestly -- progress still logs, but Stop
# won't respond mid-run.
# ---------------------------------------------------------------------------
LT_JOB = {}


def lt_job_update_ui():
    total = len(LT_JOB["work"])
    done = LT_JOB["index"]
    win_items = LT_JOB["win_items"]
    filled = int(round((done / total) * 24)) if total else 24
    try:
        win_items["LTProgText"].Text = f"Stamping lower thirds... {done} / {total}"
        win_items["LTProgBar"].Text = "\u2588" * filled + "\u2591" * (24 - filled)
    except Exception:
        pass


def lt_job_finish(stopped):
    try:
        if LT_JOB.get("timer"):
            LT_JOB["timer"].Stop()
    except Exception:
        pass
    try:
        LT_JOB["win"].Hide()
    except Exception:
        pass

    c, u, s, f = (LT_JOB["created"], LT_JOB["updated"],
                  LT_JOB["skipped"], LT_JOB["failed"])
    if stopped:
        log(f"Stopped after {LT_JOB['index']} of {len(LT_JOB['work'])} clip(s)."
            f" Already-stamped titles are kept.")
    log(f"Lower thirds done. Created {c}, updated {u}, skipped {s}, failed {f}.")
    LT_JOB.clear()


def lt_job_step(_ev=None):
    if not LT_JOB:
        return
    if LT_JOB["cancel"]:
        lt_job_finish(stopped=True)
        return

    work = LT_JOB["work"]
    if LT_JOB["index"] >= len(work):
        lt_job_finish(stopped=False)
        return

    source_name, label_text, start_frame, duration = work[LT_JOB["index"]]
    was_existing = start_frame in LT_JOB["existing_titles"]
    try:
        ok = apply_lower_third(
            LT_JOB["media_pool"], LT_JOB["existing_titles"],
            LT_JOB["template_clip"], label_text,
            start_frame, duration, LT_TITLE_VIDEO_TRACK,
        )
    except Exception:
        ok = False
        log(f"Error on {source_name}:\n{traceback.format_exc()}")

    if ok:
        if was_existing:
            LT_JOB["updated"] += 1
            log(f"Updated: {source_name} -> {label_text}")
        else:
            LT_JOB["created"] += 1
            log(f"Created: {source_name} -> {label_text}")
    else:
        LT_JOB["failed"] += 1
        log(f"Failed: {source_name} -> {label_text}")

    LT_JOB["index"] += 1
    lt_job_update_ui()


def lt_job_start(media_pool, template_clip, existing_titles, work, skipped):
    LT_JOB.clear()
    LT_JOB.update({
        "media_pool": media_pool,
        "template_clip": template_clip,
        "existing_titles": existing_titles,
        "work": work,
        "index": 0,
        "created": 0,
        "updated": 0,
        "skipped": skipped,
        "failed": 0,
        "cancel": False,
        "timer": None,
    })

    prog = disp.AddWindow(
        {
            "ID": "LTProgress",
            "WindowTitle": "Auto Lower Thirds",
            "Geometry": [220, 220, 380, 130],
            "StyleSheet": PANEL_QSS,
        },
        [
            ui.VGroup(
                {"Spacing": 8},
                [
                    ui.Label({"ID": "LTProgText", "Text": "Starting..."}),
                    ui.Label({"ID": "LTProgBar", "Text": ""}),
                    ui.Button({"ID": "BtnLTStop", "Text": "Stop"}),
                ],
            )
        ],
    )
    LT_JOB["win"] = prog
    LT_JOB["win_items"] = prog.GetItems()

    def on_stop(_ev):
        if LT_JOB:
            LT_JOB["cancel"] = True

    prog.On.BtnLTStop.Clicked = on_stop
    prog.On.LTProgress.Close = on_stop
    prog.Show()
    lt_job_update_ui()

    # Real-world testing showed the ui.Timer path creates a timer whose
    # tick events never arrive on this build -- the job window appeared
    # and nothing happened. So the job now runs synchronously: every
    # clip's progress is written to the labels and the log as it goes
    # (the log prints to the Console live regardless of repaints). The
    # Stop button is best-effort only: on builds that don't process
    # events mid-run, it won't respond until the job finishes.
    log("Processing... (progress prints to the Console live; the Stop"
        " button may not respond until done on this build)")
    while LT_JOB and LT_JOB["index"] < len(LT_JOB["work"]) and not LT_JOB["cancel"]:
        lt_job_step()
    if LT_JOB:
        lt_job_finish(stopped=bool(LT_JOB.get("cancel")))


# ---------------------------------------------------------------------------
# Button 2 -- Writer's Reel location lower thirds (ported from
# create_neighbourhood_lower_thirds.py)
# ---------------------------------------------------------------------------
def on_lower_thirds(ev):
    """Walks V1, derives each clip's neighbourhood label from its Media
    Pool folder path, and stamps or updates a lower third on V3 above it
    from the TM_LOWER_LEFT_THIRD_TEMPLATE clip. Confirms the target
    timeline before starting (this modifies the timeline and can take a
    while), then runs as a stoppable progress job."""
    if LT_JOB:
        log("Auto Lower Thirds is already running.")
        return

    project, timeline, media_pool = get_context()
    if not timeline:
        log("No timeline open -- open a project and timeline first.")
        return

    if not confirm_dialog(
        "Auto Lower Thirds",
        f"Run Auto Lower Thirds on '{timeline.GetName()}'?\n\n"
        f"This stamps a title above every labelled clip and can take a"
        f" while. It can be stopped part-way, but titles already stamped"
        f" will remain.",
    ):
        log("Auto Lower Thirds cancelled.")
        return

    try:
        resolve.OpenPage("edit")
    except Exception:
        pass

    ensure_video_track(timeline, LT_TITLE_VIDEO_TRACK)

    root_folder = media_pool.GetRootFolder()
    clip_folder_map = {}
    all_clips_by_name = {}
    walk_media_pool_folder(root_folder, [], clip_folder_map, all_clips_by_name)

    template_clip = find_template_clip_by_name(all_clips_by_name, LT_TITLE_TEMPLATE_NAME)
    if not template_clip:
        log(f"Could not find a Media Pool clip named '{LT_TITLE_TEMPLATE_NAME}'.")
        log("Create a Text+ title, drag it into the Media Pool, and name it exactly that.")
        return

    existing_titles = existing_generated_titles_by_start(timeline, LT_TITLE_VIDEO_TRACK)

    # Build the full worklist up front (fast, read-only), so the job
    # runner knows the total for the progress bar before touching
    # anything.
    work = []
    skipped = 0

    for source_track in LT_SOURCE_VIDEO_TRACKS:
        if source_track > timeline.GetTrackCount("video"):
            continue

        for source_item in timeline.GetItemListInTrack("video", source_track):
            source_name = source_item.GetName()

            try:
                media_item = source_item.GetMediaPoolItem()
            except Exception:
                media_item = None

            if not media_item:
                skipped += 1
                log(f"Skip (no Media Pool item): {source_name}")
                continue

            folder_path = clip_folder_map.get(media_pool_item_id(media_item))
            label_text = label_from_folder_path(folder_path) if folder_path else ""

            if not label_text:
                skipped += 1
                log(f"Skip (no folder label found): {source_name}")
                continue

            start_frame = timeline_item_start(source_item)
            duration = timeline_item_duration(source_item)

            if duration <= 0:
                skipped += 1
                log(f"Skip (invalid duration): {source_name}")
                continue

            if LT_DRY_RUN:
                log(f"DRY RUN -- would stamp '{label_text}' at frame {start_frame}: {source_name}")
                continue

            work.append((source_name, label_text, start_frame, duration))

    if not work:
        log(f"Nothing to do ({skipped} clip(s) skipped).")
        return

    log(f"Auto Lower Thirds: {len(work)} clip(s) to process on"
        f" '{timeline.GetName()}'.")
    lt_job_start(media_pool, template_clip, existing_titles, work, skipped)


# ---------------------------------------------------------------------------
# Timeline settings helpers (for the assemble pipeline's new timelines)
# ---------------------------------------------------------------------------
def dump_timeline_settings(target_timeline):
    """Print every setting on a timeline, sorted. Point this at a timeline
    that's already hand-configured the way you want (e.g. matching a
    delivery spec), then read off the exact key for any field you can't
    find documented, and add it to TIMELINE_SETTINGS above."""
    for key, value in sorted(target_timeline.GetSetting().items()):
        print(f"{key}: {value}")


def verify_and_fix_timeline_settings(target_timeline, settings):
    """Check each setting against its expected value and only write the
    ones that differ, logging what was already right, what got fixed,
    and what refused to change."""
    for key, value in settings.items():
        try:
            current = target_timeline.GetSetting(key)
        except Exception:
            current = None
        if str(current) == str(value):
            log(f"  OK (already set): {key} = {value}")
            continue
        ok = target_timeline.SetSetting(key, str(value))
        if ok:
            log(f"  FIXED: {key}: {current} -> {value}")
        else:
            log(f"  FAILED to set: {key} = {value} (was {current})")


# ---------------------------------------------------------------------------
# Colour section -- timeline dropdown + client grade groups
# ---------------------------------------------------------------------------
def find_timeline_by_name(project, name):
    for i in range(1, project.GetTimelineCount() + 1):
        tl = project.GetTimelineByIndex(i)
        if tl and tl.GetName() == name:
            return tl
    return None


def colour_dialog(project):
    """Blocking setup dialog for Colour Grading Prep: pick the timeline
    and the client. Returns (timeline_name, client_name) or None."""
    dlg_disp = bmd.UIDispatcher(ui)
    result = {"choice": None}

    dlg = dlg_disp.AddWindow(
        {
            "ID": "ColourDlg",
            "WindowTitle": "Colour Grading Prep",
            "Geometry": [220, 220, 420, 200],
            "StyleSheet": PANEL_QSS,
        },
        [
            ui.VGroup(
                {"Spacing": 8},
                [
                    ui.Label({"Text": "Timeline"}),
                    ui.ComboBox({"ID": "ColourTimelineCombo"}),
                    ui.Label({"Text": "Client"}),
                    ui.ComboBox({"ID": "ColourClientCombo"}),
                    ui.HGroup({"Spacing": 8, "Weight": 0}, [
                        ui.Button({"ID": "BtnColourCancel", "Text": "Cancel"}),
                        ui.Button({"ID": "BtnColourRun", "Text": "Run"}),
                    ]),
                ],
            )
        ],
    )

    ditems = dlg.GetItems()
    for i in range(1, project.GetTimelineCount() + 1):
        tl = project.GetTimelineByIndex(i)
        if tl:
            ditems["ColourTimelineCombo"].AddItem(tl.GetName())
    for client_name in CLIENT_COLOR_PRESETS:
        ditems["ColourClientCombo"].AddItem(client_name)

    def on_run(_ev):
        tl_name = ditems["ColourTimelineCombo"].CurrentText
        client = ditems["ColourClientCombo"].CurrentText
        if tl_name and client:
            result["choice"] = (tl_name, client)
        dlg_disp.ExitLoop()

    def on_cancel(_ev):
        dlg_disp.ExitLoop()

    dlg.On.BtnColourRun.Clicked = guard(on_run)
    dlg.On.BtnColourCancel.Clicked = on_cancel
    dlg.On.ColourDlg.Close = on_cancel

    hold_log_widget()
    dlg.Show()
    run_loop_resilient(dlg_disp, "colour dialog")
    dlg.Hide()
    release_log_widget()
    return result["choice"]


def on_apply_client_color(ev):
    """Assign every video clip on the chosen timeline to the chosen
    client's colour group (creating the group if needed), apply the
    client's clip-level .drx grade if one is configured, and verify/fix
    the timeline's format + colour settings."""
    project, _, _ = get_context()
    if not project:
        log("No project open.")
        return

    choice = colour_dialog(project)
    if not choice:
        log("Colour Grading Prep cancelled.")
        return
    timeline_name, client_name = choice

    preset = CLIENT_COLOR_PRESETS.get(client_name)
    if not preset:
        log(f"Unknown client: {client_name}")
        return

    target = find_timeline_by_name(project, timeline_name)
    if not target:
        log(f"Timeline '{timeline_name}' not found.")
        return

    project.SetCurrentTimeline(target)
    log(f"Applying {client_name} colour setup to '{timeline_name}'...")

    # --- Find or create the colour group -------------------------------
    group_name = preset["group_name"]
    group = None
    for existing in project.GetColorGroupsList() or []:
        try:
            if existing.GetName() == group_name:
                group = existing
                break
        except Exception:
            pass

    if group is None:
        group = project.AddColorGroup(group_name)
        if group:
            log(f"Created colour group '{group_name}'. NOTE: its pre/post-clip"
                f" group grades are empty -- they need their one-time manual"
                f" setup (the API can't author group node trees).")
        else:
            log(f"Could not create colour group '{group_name}'.")
            return
    else:
        log(f"Using existing colour group '{group_name}' (pre/post grades"
            f" already on it will apply to everything assigned).")

    # --- Assign every video clip on every track -------------------------
    assigned = skipped = 0
    clips_for_grade = []
    for track_index in range(1, target.GetTrackCount("video") + 1):
        for item in target.GetItemListInTrack("video", track_index) or []:
            name = item.GetName() or ""
            # Generated lower-third titles aren't camera footage --
            # keep them out of the client grade group.
            if name.startswith(LT_GENERATED_TITLE_PREFIX):
                skipped += 1
                continue
            try:
                if item.AssignToColorGroup(group):
                    assigned += 1
                    clips_for_grade.append(item)
                else:
                    skipped += 1
                    log(f"  Could not assign: {name}")
            except Exception:
                skipped += 1
                log(f"  Could not assign: {name}")

    log(f"Assigned {assigned} clip(s) to '{group_name}' ({skipped} skipped).")

    # --- Apply the clip-level grade from the client's .drx still --------
    drx_path = preset.get("drx_path") or ""
    if drx_path and clips_for_grade:
        try:
            ok = target.ApplyGradeFromDRX(drx_path, 0, clips_for_grade)
            if ok:
                log(f"Applied clip grade from {drx_path}")
            else:
                log(f"ApplyGradeFromDRX failed for {drx_path} -- check the"
                    f" path and that the .drx was exported from a still.")
        except Exception:
            log(f"Error applying .drx grade:\n{traceback.format_exc()}")
    elif not drx_path:
        log(f"No .drx configured for {client_name} -- skipped clip-grade"
            f" step. Set drx_path in CLIENT_COLOR_PRESETS to enable it.")

    # --- Verify / fix timeline + colour settings ------------------------
    log("Checking timeline settings...")
    verify_and_fix_timeline_settings(target, TIMELINE_SETTINGS)
    log(f"{client_name} colour setup done.")


# ---------------------------------------------------------------------------
# Mid/Short Form Assembly -- script parsing (validated against the real
# Skyscanner Seoul docx and Expedia Amsterdam docx export; see
# extract_pois.py for the standalone, testable version of this logic)
# ---------------------------------------------------------------------------
SEGMENT_RE = re.compile(
    r"^\s*(?:(?:Segment|Theme)\s*\d+\s*:\s*(?P<name>.+?)|(?P<marker>INTRO|OUTRO))"
    r"\s*(?:\(\s*\d+\s*words?\s*\))?\s*$",
    re.IGNORECASE,
)


def _iter_doc_paragraphs(doc):
    """Yield every paragraph in document order, including inside tables --
    Pages-exported scripts keep the whole script body in a table."""
    from docx.document import Document as _DocumentBody
    from docx.oxml.ns import qn
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    def walk(parent):
        parent_elm = parent.element.body if isinstance(parent, _DocumentBody) else parent._tc
        for child in parent_elm.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                for row in Table(child, parent).rows:
                    for cell in row.cells:
                        yield from walk(cell)

    yield from walk(doc)


TIME_PREFIX_RE = re.compile(
    r"^\s*\d{1,2}([:.]\d{2})?\s*(am|pm)\b[\s,\-\u2013]*", re.IGNORECASE
)


def _clean_segment_name(name):
    """Strip a leading time like '10:30am, ' from a segment name, so the
    group label and lower third read 'Jordaan and De Negen Straatjes'
    rather than the schedule entry."""
    cleaned = TIME_PREFIX_RE.sub("", name).strip()
    return cleaned or name.strip()


def extract_script_structure(docx_path):
    """Parse a writer's script into ordered segments with bolded POIs and
    plain narration. Requires python-docx installed into Resolve's Python.

    Header detection is deliberately paranoid, because real scripts vary:
      - a header as its own paragraph (bold or plain)          -> caught
      - several header lines soft-broken in one bold paragraph -> caught
      - a bold header run sharing a paragraph with narration   -> caught
    """
    from docx import Document

    doc = Document(docx_path)
    segments = []
    current = {"segment": None, "pois": [], "narration": ""}

    def push():
        if current["segment"] or current["pois"] or current["narration"].strip():
            current["narration"] = current["narration"].strip()
            segments.append(dict(current))

    def start_segment(m):
        nonlocal_current = m.group("name") or m.group("marker").upper()
        push()
        current.clear()
        current.update({"segment": _clean_segment_name(nonlocal_current),
                        "pois": [], "narration": ""})

    for para in _iter_doc_paragraphs(doc):
        text = para.text.strip()
        if not text:
            continue

        # Whole paragraph is a single header line (bold or plain)
        m = SEGMENT_RE.match(text)
        if m:
            start_segment(m)
            continue

        # Pop-up annotation paragraphs: bold runs are annotation content
        if any(r.bold and r.italic for r in para.runs):
            continue

        has_plain = any((not r.bold) and r.text.strip() for r in para.runs)

        if not has_plain:
            # Fully-bold paragraph: may hold several soft-broken header
            # lines; anything else in it is a section header, not a POI.
            for line in text.splitlines():
                lm = SEGMENT_RE.match(line.strip())
                if lm:
                    start_segment(lm)
            continue

        # Mixed paragraph: walk runs in order, so a bold header run that
        # shares a paragraph with narration still splits the segment at
        # the right point.
        for run in para.runs:
            rtext = run.text
            if not rtext:
                continue
            if run.bold and not run.italic:
                for piece in rtext.split("\n"):
                    piece = piece.strip()
                    if not piece:
                        continue
                    pm = SEGMENT_RE.match(piece)
                    if pm:
                        start_segment(pm)
                        continue
                    t = piece.strip("[](),.:;\u2014-").strip()
                    if len(t) > 1 and t.lower() not in (p.lower() for p in current["pois"]):
                        current["pois"].append(t)
            else:
                current["narration"] += rtext
        current["narration"] += "\n"

    push()
    return segments


def _fold_accents(text):
    """Fold accents so non-English or accent-drifted names match their
    ascii-typed counterparts (Menilmontant matches the accented form)."""
    try:
        return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    except Exception:
        return text


def _norm_name(name):
    n = _fold_accents(name).strip().lower()
    if n.endswith("(n)"):
        n = n[:-3].strip()
    n = re.sub(r"[\u2019'\-\u2013\u2014]", " ", n)
    n = re.sub(r"\s+", " ", n)
    return n.strip()


def poi_matches_bin(poi_name, bin_name):
    """Assembly matching now uses the same tuned, tiered fuzzy matcher
    as Sort by Shoot Notes -- exact / containment / misspelling-tolerant
    / distinctive-token -- with accent folding for non-English names."""
    return _sort_match_tier(poi_name, bin_name) > 0


# ---------------------------------------------------------------------------
# Mid/Short Form Assembly -- matching + build
# ---------------------------------------------------------------------------
def clips_matching_location(location_name, clip_folder_map):
    """All media pool clip ids whose bin path contains a folder matching
    the location name, at any depth."""
    matched = set()
    for clip_id, folder_path in clip_folder_map.items():
        for folder_name in folder_path:
            if poi_matches_bin(location_name, folder_name):
                matched.add(clip_id)
                break
    return matched


def build_location_groups(structure, clip_folder_map, package_order):
    """Turn the parsed script into ordered groups:

        (label, [package entries], [extra pool clip ids])

    Matching per segment, all filtered/ordered against the Clip Asset
    Package (the approved cut):
      1. segment name matches bins (neighbourhood folders)
      2. bolded POIs match bins (Skyscanner style)
      3. NARRATION mentions bin names (covers segments whose locations
         only exist as sub-location bins -- e.g. a "Centrum" segment
         whose clips live in "Dam Square" and "Damrak" bins)
    "Extras" are pool clips under the same matched locations that were
    never used in the package -- appended at the end of each group.
    A clip only joins the first group that claims it.
    """
    package_ids = [e["cid"] for e in package_order]
    package_id_set = set(package_ids)
    groups = []
    misses = []
    used = set()        # package_order indexes already claimed
    extras_used = set() # pool clip ids already claimed as extras

    # Every folder name in the pool, for narration mention-matching
    all_folder_names = set()
    for folder_path in clip_folder_map.values():
        all_folder_names.update(folder_path)
    GENERIC_FOLDERS = {"master", "video", "footage", "clips", "media"}

    def guarded(label, ids):
        """Refuse absurdly broad matches (e.g. a city-level bin that
        holds most of the package) so one segment can't swallow all."""
        in_package = len(ids & package_id_set)
        if len(structure) > 2 and in_package > max(4, int(len(package_ids) * 0.6)):
            log(f"  (ignored over-broad match for '{label}' --"
                f" {in_package} of {len(package_ids)} package clips)")
            return set()
        return ids

    def narration_matches(seg):
        matched = set()
        if not seg["narration"] or seg["segment"] in (None, "INTRO", "OUTRO"):
            return matched
        narr = _sort_norm(seg["narration"])
        narr_words = narr.split()
        for folder_name in all_folder_names:
            nb = _sort_norm(folder_name)
            if len(nb) < 4 or nb in GENERIC_FOLDERS:
                continue
            hit = bool(re.search(r"\b" + re.escape(nb) + r"\b", narr))
            if not hit:
                # Fuzzy sliding window over the narration: a slightly
                # misspelled or accent-drifted folder name still counts
                # as mentioned. Also inherently reaches folders OUTSIDE
                # the neighbourhood stack, anywhere in the pool.
                k = max(1, len(nb.split()))
                for i in range(0, max(0, len(narr_words) - k + 1)):
                    window = " ".join(narr_words[i:i + k])
                    if abs(len(window) - len(nb)) <= 3 and \
                       SequenceMatcher(None, window, nb).ratio() >= 0.86:
                        hit = True
                        break
            if hit:
                ids = {cid for cid, fp in clip_folder_map.items()
                       if folder_name in fp}
                matched |= guarded(folder_name, ids)
        return matched

    def add_group(label, matched_ids):
        entries = [entry for idx, entry in enumerate(package_order)
                   if entry["cid"] in matched_ids and idx not in used]
        extras = [cid for cid in clip_folder_map
                  if cid in matched_ids
                  and cid not in package_id_set
                  and cid not in extras_used]
        if entries or extras:
            used.update(idx for idx, entry in enumerate(package_order)
                        if entry["cid"] in matched_ids and idx not in used)
            extras_used.update(extras)
            groups.append((label, entries, extras))
            return True
        return False

    for seg in structure:
        seg_name = seg["segment"]
        matched = set()
        if seg_name:
            matched |= guarded(seg_name,
                               clips_matching_location(seg_name, clip_folder_map))
        matched |= narration_matches(seg)

        seg_grouped = False
        if matched and seg_name and seg_name not in ("INTRO", "OUTRO"):
            seg_grouped = add_group(seg_name, matched)
        if seg_name and seg_name not in ("INTRO", "OUTRO") and not seg_grouped:
            misses.append(seg_name)

        for poi in seg["pois"]:
            poi_matched = guarded(poi, clips_matching_location(poi, clip_folder_map))
            if poi_matched:
                if not add_group(poi, poi_matched):
                    pass  # everything already claimed by earlier groups
            else:
                misses.append(poi)

    return groups, misses


AUDIO_EXTENSIONS = (".wav", ".mp3", ".aif", ".aiff", ".m4a", ".flac", ".ogg")
GRAPHIC_EXTENSIONS = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".psd",
                      ".exr", ".tga", ".bmp", ".gif", ".svg", ".ai")


def is_excluded_extra(media_pool_item):
    """Extras must be real footage: no music/VO, no stills/graphics, no
    titles or generators."""
    if is_audio_clip(media_pool_item):
        return True
    try:
        clip_type = (media_pool_item.GetClipProperty("Type") or "").lower()
        if any(word in clip_type for word in ("still", "graphic", "title",
                                              "generator", "matte")):
            return True
    except Exception:
        pass
    try:
        name = (media_pool_item.GetName() or "").lower()
        if name.endswith(GRAPHIC_EXTENSIONS):
            return True
    except Exception:
        pass
    return False


def is_audio_clip(media_pool_item):
    """Music/VO files living in POI bins must never be pulled into an
    assembly as extras."""
    try:
        clip_type = (media_pool_item.GetClipProperty("Type") or "").lower()
        if "audio" in clip_type and "video" not in clip_type:
            return True
    except Exception:
        pass
    try:
        name = (media_pool_item.GetName() or "").lower()
        if name.endswith(AUDIO_EXTENSIONS):
            return True
    except Exception:
        pass
    return False


def group_neighbourhood(entry_cids, clip_folder_map):
    """The neighbourhood a group's clips live in: the (N)-marked ancestor
    folder that most of the group's clips share. Returns the folder name
    with the (N) suffix stripped, or None when the clips have no (N)
    ancestor."""
    counts = {}
    for cid in entry_cids:
        for folder in (clip_folder_map.get(cid) or []):
            stripped = folder.strip()
            if stripped.lower().endswith("(n)"):
                name = stripped[:-3].strip()
                counts[name] = counts.get(name, 0) + 1
                break
    if not counts:
        return None
    return max(counts, key=counts.get)


def lower_third_label(script_label, neighbourhood):
    """'POI NAME - NEIGHBOURHOOD', except when that would duplicate --
    if the script label already IS (or contains) the neighbourhood, the
    label stands alone."""
    if not neighbourhood:
        return script_label
    if _sort_norm(neighbourhood) in _sort_norm(script_label):
        return script_label
    return f"{script_label} - {neighbourhood}"


def clip_frames(media_pool_item):
    try:
        return int(media_pool_item.GetClipProperty("Frames") or 0)
    except Exception:
        return 0


def unique_timeline_name(project, base_name):
    existing = {project.GetTimelineByIndex(i).GetName()
                for i in range(1, project.GetTimelineCount() + 1)
                if project.GetTimelineByIndex(i)}
    if base_name not in existing:
        return base_name
    n = 2
    while f"{base_name} {n}" in existing:
        n += 1
    return f"{base_name} {n}"


def run_assembly(project, media_pool, params):
    # --- Parse -----------------------------------------------------------
    try:
        structure = extract_script_structure(params["script_path"])
    except ImportError:
        log("python-docx is not installed in Resolve's Python.")
        log("In Resolve's Console run:  import sys; print(sys.executable)")
        log("Then in Terminal:  <that python> -m pip install python-docx")
        return
    except Exception:
        log(f"Could not parse the script:\n{traceback.format_exc()}")
        return

    n_pois = sum(len(s["pois"]) for s in structure)
    log(f"Parsed script: {len(structure)} segment(s), {n_pois} bolded POI(s).")

    # --- Index the media pool ---------------------------------------------
    clip_folder_map = {}
    all_clips_by_name = {}
    clip_by_id = {}
    walk_media_pool_folder(media_pool.GetRootFolder(), [], clip_folder_map,
                            all_clips_by_name, clip_by_id)

    # --- Collect the Clip Asset Package's clips, in timeline order --------
    package_tl = find_timeline_by_name(project, params["package_name"])
    if not package_tl:
        log(f"Timeline '{params['package_name']}' not found.")
        return

    # Each PACKAGE TIMELINE ITEM is one assembly clip -- carrying its own
    # in/out trims, grade, and colour group. The media pool clip is only
    # used for bin-path location matching and as the append source.
    package_order = []
    seen_usage = set()
    for track_index in range(1, package_tl.GetTrackCount("video") + 1):
        for item in package_tl.GetItemListInTrack("video", track_index) or []:
            name = item.GetName() or ""
            if name.startswith(LT_GENERATED_TITLE_PREFIX):
                continue
            try:
                mpi = item.GetMediaPoolItem()
            except Exception:
                mpi = None
            if not mpi:
                continue
            cid = media_pool_item_id(mpi)
            if not cid:
                continue
            try:
                src_start = int(item.GetSourceStartFrame())
                # Exclusive end from the item's true duration -- using
                # GetSourceEndFrame() loses the final frame.
                src_end = src_start + int(item.GetDuration(False))
            except Exception:
                src_start = src_end = 0
            usage_key = (cid, src_start, src_end)
            if usage_key in seen_usage:
                continue  # identical duplicate usage
            seen_usage.add(usage_key)
            package_order.append({
                "cid": cid,
                "mpi": mpi,
                "item": item,
                "start": src_start,
                "end": src_end,
            })

    log(f"Clip Asset Package '{params['package_name']}':"
        f" {len(package_order)} clip(s).")
    if not package_order:
        log("Nothing to assemble.")
        return

    # --- Match into location groups ---------------------------------------
    groups, misses = build_location_groups(structure, clip_folder_map, package_order)
    if misses:
        log(f"  Script locations with no matching clips:"
            f" {', '.join(misses)}")
    if not groups:
        log("No package clips matched any script location -- check that bin"
            " names and script POIs line up.")
        return

    for label, entries, extras in groups:
        log(f"  {label}: {len(entries)} package clip(s)"
            f" + {len(extras)} unused bin clip(s)")
    unmatched = len(package_order) - sum(len(e) for _, e, _ in groups)
    if unmatched:
        log(f"  ({unmatched} package clip(s) matched no script location --"
            f" they'll be placed together at the end)")

    # --- Create + configure the new timeline ------------------------------
    tl_name = unique_timeline_name(project, params["new_name"])
    new_tl = media_pool.CreateEmptyTimeline(tl_name)
    if not new_tl:
        log(f"Could not create timeline '{tl_name}'.")
        return
    project.SetCurrentTimeline(new_tl)
    log(f"Created '{tl_name}'. Applying timeline settings...")
    verify_and_fix_timeline_settings(new_tl, TIMELINE_SETTINGS)

    try:
        fps = float(new_tl.GetSetting("timelineFrameRate") or 29.97)
    except Exception:
        fps = 29.97
    gap_frames = int(round(params["gap_seconds"] * fps))
    mode = params.get("mode", "lisa")  # lisa: package + bin extras;
                                       # maggie: package clips only

    # --- Lower thirds template (optional but expected) ---------------------
    template_clip = find_template_clip_by_name(all_clips_by_name, LT_TITLE_TEMPLATE_NAME)
    if template_clip:
        ensure_video_track(new_tl, LT_TITLE_VIDEO_TRACK)
    else:
        log(f"No '{LT_TITLE_TEMPLATE_NAME}' clip in the Media Pool -- groups"
            f" will be assembled without lower thirds.")

    # --- Append groups with gaps, lower third per group --------------------
    cursor = int(new_tl.GetStartFrame())
    existing_titles = {}
    placed = failed = grade_fails = 0

    def place_package_entry(entry, record_frame):
        """Append one Clip Asset Package entry with its exact trims,
        grade, and colour group. Returns frames placed (0 on failure)."""
        nonlocal placed, failed, grade_fails
        mpi = entry["mpi"]
        frames = entry["end"] - entry["start"]
        if frames <= 0:
            # No usable trim info -- fall back to full clip length
            frames = clip_frames(mpi)
            entry = dict(entry, start=None, end=None)
        if frames <= 0:
            failed += 1
            log(f"  Skipped (no frame count): {mpi.GetName()}")
            return 0

        clip_info = {
            "mediaPoolItem": mpi,
            "trackIndex": 1,
            "recordFrame": record_frame,
        }
        if entry["start"] is not None:
            # Carry the package edit's exact in/out points
            clip_info["startFrame"] = entry["start"]
            clip_info["endFrame"] = entry["end"]

        appended = media_pool.AppendToTimeline([clip_info]) or []
        if not appended:
            failed += 1
            log(f"  Failed to append: {mpi.GetName()}")
            return 0

        placed += 1
        new_item = appended[0]

        # Copy the package item's clip grade + colour group across, so
        # the assembly matches the approved look -- not the raw bin
        # version of the clip.
        try:
            if not entry["item"].CopyGrades([new_item]):
                grade_fails += 1
        except Exception:
            grade_fails += 1
        try:
            group = entry["item"].GetColorGroup()
            if group:
                new_item.AssignToColorGroup(group)
        except Exception:
            pass
        return frames

    for label, entries, extras in groups:
        group_start = cursor
        group_frames = 0

        for entry in entries:
            frames = place_package_entry(entry, cursor)
            if frames:
                cursor += frames
                group_frames += frames

        # Unused bin clips for this location, after the approved package
        # clips -- rough alternates for the editor. A short divider gap
        # separates the approved material from the leftovers, and each
        # extra is forced to FULL length (explicit start/end frames), so
        # any in/out marks left on the pool clips are ignored.
        if mode != "lisa":
            extras = []  # Maggie: package clips only, no bin remainder

        if extras and group_frames > 0:
            divider = int(round(ASSEMBLY_EXTRAS_GAP_SECONDS * fps))
            cursor += divider
            group_frames += divider  # keep the lower third spanning the lot

        for cid in extras:
            mpi = clip_by_id.get(cid)
            if not mpi:
                continue
            if is_excluded_extra(mpi):
                continue  # no music/VO, stills, graphics, or titles
            frames = clip_frames(mpi)
            if frames <= 0:
                continue
            appended = media_pool.AppendToTimeline([{
                "mediaPoolItem": mpi,
                "trackIndex": 1,
                "recordFrame": cursor,
                "startFrame": 0,        # ignore any in/out marks on the
                "endFrame": frames,     # pool clip -- always full length
            }]) or []
            if appended:
                placed += 1
                cursor += frames
                group_frames += frames
            else:
                failed += 1
                log(f"  Failed to append extra: {mpi.GetName()}")

        if group_frames > 0 and template_clip:
            nei = group_neighbourhood(
                [e["cid"] for e in entries] + list(extras), clip_folder_map)
            lt_text = lower_third_label(label, nei)
            ok = apply_lower_third(media_pool, existing_titles, template_clip,
                                    lt_text, group_start, group_frames,
                                    LT_TITLE_VIDEO_TRACK)
            log(f"  -> lower third: {label}" + ("" if ok else " (failed)"))

        if group_frames > 0:
            cursor += gap_frames

    # Everything from the Clip Asset Package the script never claimed,
    # laid together at the very end (both Lisa and Maggie).
    used_ids = {id(e) for _, es, _ in groups for e in es}
    leftover = [e for e in package_order if id(e) not in used_ids]
    if leftover:
        tail_start = cursor
        tail_frames = 0
        for entry in leftover:
            frames = place_package_entry(entry, cursor)
            if frames:
                cursor += frames
                tail_frames += frames
        if tail_frames > 0 and template_clip:
            apply_lower_third(media_pool, existing_titles, template_clip,
                              "Unused Clip Assets", tail_start, tail_frames,
                              LT_TITLE_VIDEO_TRACK)
        log(f"  Unused package clips placed at the end: {len(leftover)}")

    log(f"Assembly done. Placed {placed} clip(s) in {len(groups)} group(s)"
        f" on '{tl_name}' ({failed} failed).")
    if grade_fails:
        log(f"  Note: grade copy failed on {grade_fails} clip(s) -- those"
            f" carry the bin-version look. (CopyGrades needs Resolve 18.5+.)")


# ---------------------------------------------------------------------------
# Mid/Short Form Assembly -- setup dialog
# ---------------------------------------------------------------------------
def assemble_dialog(project):
    """Blocking setup dialog: script file, asset package timeline, form,
    new timeline name, gap. Returns a params dict, or None on cancel."""
    dlg_disp = bmd.UIDispatcher(ui)
    result = {"params": None}

    timeline_names = []
    for i in range(1, project.GetTimelineCount() + 1):
        tl = project.GetTimelineByIndex(i)
        if tl:
            timeline_names.append(tl.GetName())

    dlg = dlg_disp.AddWindow(
        {
            "ID": "AssembleDlg",
            "WindowTitle": "Mid/Short Form Assembly",
            "Geometry": [180, 180, 480, 330],
            "StyleSheet": PANEL_QSS,
        },
        [
            ui.VGroup(
                {"Spacing": 8},
                [
                    ui.Label({"Text": "Script (.docx)"}),
                    ui.HGroup({"Spacing": 6, "Weight": 0}, [
                        ui.LineEdit({"ID": "ScriptPath",
                                     "PlaceholderText": "Path to the writer's script"}),
                        ui.Button({"ID": "BtnBrowse", "Text": "Browse", "Weight": 0}),
                    ]),
                    ui.Label({"Text": "Asset package timeline"}),
                    ui.ComboBox({"ID": "PackageCombo"}),
                    ui.Label({"Text": "Form"}),
                    ui.ComboBox({"ID": "FormCombo"}),
                    ui.Label({"Text": "New timeline name"}),
                    ui.LineEdit({"ID": "NewName",
                                 "PlaceholderText": "Leave empty for automatic name"}),
                    ui.Label({"Text": "Version"}),
                    ui.ComboBox({"ID": "VersionCombo"}),
                    ui.Label({"Text": "Gap between groups"}),
                    ui.ComboBox({"ID": "GapCombo"}),
                    ui.Label({"ID": "DlgWarn", "Text": ""}),
                    ui.HGroup({"Spacing": 8, "Weight": 0}, [
                        ui.Button({"ID": "BtnAsmCancel", "Text": "Cancel"}),
                        ui.Button({"ID": "BtnAsmRun", "Text": "Run assembly"}),
                    ]),
                ],
            )
        ],
    )

    ditems = dlg.GetItems()
    for name in timeline_names:
        ditems["PackageCombo"].AddItem(name)
    for form in ("Long Form", "Mid Form", "Short Form"):
        ditems["FormCombo"].AddItem(form)
    ditems["FormCombo"].CurrentIndex = 1
    for gap in ("5 seconds", "15 seconds", "30 seconds"):
        ditems["GapCombo"].AddItem(gap)
    ditems["GapCombo"].CurrentIndex = 1
    ditems["VersionCombo"].AddItem("Lisa -- package + rest of location bins")
    ditems["VersionCombo"].AddItem("Maggie -- package clips only (faster)")

    def on_browse(_ev):
        try:
            path = fusion.RequestFile()
            if path:
                ditems["ScriptPath"].Text = str(path)
        except Exception:
            ditems["DlgWarn"].Text = "File browser unavailable -- paste the path."

    def on_run(_ev):
        script_path = (ditems["ScriptPath"].Text or "").strip()
        if not script_path or not os.path.isfile(script_path):
            ditems["DlgWarn"].Text = "Script file not found -- check the path."
            return
        if not script_path.lower().endswith(".docx"):
            ditems["DlgWarn"].Text = ("Scripts must be .docx (export from Pages"
                                       " or Google Docs first).")
            return
        package_name = ditems["PackageCombo"].CurrentText
        if not package_name:
            ditems["DlgWarn"].Text = "No asset package timeline selected."
            return
        form = ditems["FormCombo"].CurrentText
        new_name = (ditems["NewName"].Text or "").strip()
        if not new_name:
            new_name = f"{package_name} - {form} Assembly"
        gap_seconds = int((ditems["GapCombo"].CurrentText or "15 seconds").split()[0])
        version_text = ditems["VersionCombo"].CurrentText or ""
        result["params"] = {
            "mode": "maggie" if version_text.startswith("Maggie") else "lisa",
            "script_path": script_path,
            "package_name": package_name,
            "form": form,
            "new_name": new_name,
            "gap_seconds": gap_seconds,
        }
        dlg_disp.ExitLoop()

    def on_cancel(_ev):
        dlg_disp.ExitLoop()

    dlg.On.BtnBrowse.Clicked = guard(on_browse)
    dlg.On.BtnAsmRun.Clicked = guard(on_run)
    dlg.On.BtnAsmCancel.Clicked = on_cancel
    dlg.On.AssembleDlg.Close = on_cancel

    hold_log_widget()
    dlg.Show()
    run_loop_resilient(dlg_disp, "setup dialog")
    dlg.Hide()
    release_log_widget()
    return result["params"]


# ---------------------------------------------------------------------------
# Button 3 -- Mid/Short Form Assembly
# ---------------------------------------------------------------------------
def on_assemble(ev):
    """Full pipeline: parse the script's segments/POIs, match them to
    media pool bins, filter to the Clip Asset Package's clips, and build
    a new configured timeline of location groups with gaps and a lower
    third per group."""
    project, timeline, media_pool = get_context()
    if not project:
        log("No project open.")
        return

    params = assemble_dialog(project)
    if not params:
        log("Assembly cancelled.")
        return

    log(f"Assembling '{params['new_name']}' ({params['form']}) from"
        f" '{params['package_name']}'...")
    run_assembly(project, media_pool, params)


# ---------------------------------------------------------------------------
# Button -- rename video & audio tracks to the standard layout
# ---------------------------------------------------------------------------
def on_rename_tracks(ev):
    """Apply the standard track layout to the current timeline. Renames
    existing tracks and creates any that are missing, so the full layout
    (V1-V3, A1-A7) always ends up present."""
    project, timeline, media_pool = get_context()
    if not timeline:
        log("No timeline open -- open a project and timeline first.")
        return

    log(f"Applying standard track layout to '{timeline.GetName()}'...")

    renamed = created = failed = 0

    # --- Video tracks ----------------------------------------------------
    for index in sorted(VIDEO_TRACK_NAMES):
        name = VIDEO_TRACK_NAMES[index]
        while timeline.GetTrackCount("video") < index:
            if not timeline.AddTrack("video"):
                log(f"  FAILED to create video track V{index}")
                failed += 1
                break
            created += 1
        if timeline.GetTrackCount("video") < index:
            continue

        current = None
        try:
            current = timeline.GetTrackName("video", index)
        except Exception:
            pass

        if current == name:
            log(f"  OK (already named): V{index} = {name}")
        elif timeline.SetTrackName("video", index, name):
            renamed += 1
            log(f"  V{index}: {current or '(unnamed)'} -> {name}")
        else:
            failed += 1
            log(f"  FAILED to rename V{index} to {name}")

    # --- Audio tracks ----------------------------------------------------
    for index in sorted(AUDIO_TRACK_NAMES):
        name, subtype = AUDIO_TRACK_NAMES[index]
        while timeline.GetTrackCount("audio") < index:
            if not timeline.AddTrack("audio", subtype):
                log(f"  FAILED to create audio track A{index}")
                failed += 1
                break
            created += 1
        if timeline.GetTrackCount("audio") < index:
            continue

        current = None
        try:
            current = timeline.GetTrackName("audio", index)
        except Exception:
            pass

        if current == name:
            log(f"  OK (already named): A{index} = {name}")
        elif timeline.SetTrackName("audio", index, name):
            renamed += 1
            log(f"  A{index}: {current or '(unnamed)'} -> {name}")
        else:
            failed += 1
            log(f"  FAILED to rename A{index} to {name}")

    log(f"Track layout done. Renamed {renamed}, created {created}, failed {failed}.")


# ---------------------------------------------------------------------------
# Sort by Shoot Notes -- parse a shoot-notes docx into themes/POIs and
# reorganise the destination's POI bins into theme/POI folders, colouring
# clips per theme. Validated against the real London shoot notes and the
# real London bin names (including misspellings like "Greenwhich").
# ---------------------------------------------------------------------------
SORT_PARENT_FOLDER = "RAW-FILES"   # theme folders are created in here
SORT_EXTRAS_NAME = "EXTRAS"
SORT_FOLDER_COLORS = ["Blue", "Cyan", "Green", "Yellow",
                      "Orange", "Red", "Pink", "Purple"]
# Clip colours are a different palette to bin labels -- no Cyan/Red, so
# those themes use the nearest clip colour.
SORT_CLIP_COLOR = {"Blue": "Blue", "Cyan": "Teal", "Green": "Green",
                   "Yellow": "Yellow", "Orange": "Orange",
                   "Red": "Chocolate", "Pink": "Pink", "Purple": "Purple"}
# Folders that are never location bins -- straight to EXTRAS.
SORT_FOLDER_STOPLIST = {"general", "misc", "extras", "assets", "temp"}

SORT_THEME_RE = re.compile(r"^Theme\s*(\d+)\s*:\s*(.+)$", re.IGNORECASE)
SORT_POI_LABEL_RE = re.compile(r"^(POI\s*\d+\.\d+|Contingency POI)$", re.IGNORECASE)
SORT_GENERIC_WORDS = {
    "palace", "park", "parks", "museum", "market", "street", "st", "road",
    "gardens", "garden", "gallery", "house", "bridge", "square", "centre",
    "center", "hill", "lane", "town", "city", "district", "abbey", "london",
    "the", "and", "of", "at",
    # directional/civic modifiers -- shared by unrelated places ("National
    # Gallery" vs "National Theatre"), so never distinctive on their own
    "national", "royal", "east", "west", "north", "south", "central", "greater",
}


def _sort_norm(name):
    n = _fold_accents(name).strip().lower()
    for suffix in ("(n)", "(r)", "(c)"):
        if n.endswith(suffix):
            n = n[: -len(suffix)].strip()
    n = re.sub(r"[\u2019']", "", n)   # apostrophes deleted: st paul's -> st pauls
    n = re.sub(r"[\-\u2013\u2014!_]", " ", n)
    n = re.sub(r"\b\d+\b", " ", n)  # drop numeric ID suffixes like 44304425
    n = re.sub(r"\s+", " ", n)
    return n.strip()


def _sort_title_locations(title):
    """'London Docklands (Canary Wharf/West India Quay)' -> the separate
    location names inside a POI title."""
    parts = []
    m = re.match(r"^(.*?)\((.*?)\)\s*$", title)
    inner = ""
    if m:
        title, inner = m.group(1).strip(), m.group(2)
    for chunk in re.split(r"\s+and\s+|/|,", title):
        chunk = chunk.strip()
        if chunk:
            parts.append(chunk)
    for chunk in re.split(r"\s+and\s+|/|,", inner):
        chunk = chunk.strip()
        if chunk and not chunk.lower().startswith(("murals", "street art")):
            parts.append(chunk)
    return parts


def _sort_candidate_locations(lines):
    """Location candidates from shot-suggestion bullets: 'Brick Lane:
    Capture...' prefixes and short plain lines like 'Ebor Street'.
    Over-generation is harmless -- candidates only matter if they match a
    real bin name -- but 'General...' shot descriptions are dropped so a
    'General' bin can never be claimed by a theme."""
    out = []
    for line in lines:
        line = line.strip().lstrip("*\u2022-\u2013 ").strip()
        if not line:
            continue
        low = line.lower()
        if low.startswith(("shot suggestion", "note", "other notable", "general")):
            continue
        if ":" in line:
            prefix = line.split(":", 1)[0].strip()
            if 1 <= len(prefix.split()) <= 5 and prefix[:1].isupper():
                out.append(prefix)
        else:
            head = re.split(r"\s[-\u2013\u2014]\s|\(", line)[0].strip()
            if head and 1 <= len(head.split()) <= 5 and head[:1].isupper() \
                    and not head.endswith("."):
                out.append(head)
    return out


def parse_shoot_notes(docx_path):
    """-> (destination, [{"num", "name", "pois": [{"label", "title",
    "candidates"}]}]) from the labelled two-column shoot-notes table."""
    from docx import Document

    doc = Document(docx_path)
    destination = None
    themes = []
    current_theme = None

    for table in doc.tables:
        for row in table.rows:
            c0 = row.cells[0].text.strip()
            c1 = row.cells[1].text.strip()
            if c0.lower().startswith("destination:") and not destination:
                destination = c0.split(":", 1)[1].strip()
            first_line = c0.splitlines()[0].strip() if c0 else ""

            tm = SORT_THEME_RE.match(first_line)
            if tm and c0 == c1:
                current_theme = {"num": int(tm.group(1)),
                                 "name": tm.group(2).strip(), "pois": []}
                themes.append(current_theme)
                continue

            pm = SORT_POI_LABEL_RE.match(first_line)
            if pm and current_theme is not None and c1:
                lines = [l for l in c1.splitlines() if l.strip()]
                title = lines[0].strip()
                cands = (_sort_title_locations(title)
                         + _sort_candidate_locations(lines[1:]))
                seen, ordered = set(), []
                for cand in cands:
                    if cand.lower() not in seen:
                        seen.add(cand.lower())
                        ordered.append(cand)
                current_theme["pois"].append({
                    "label": first_line,
                    "title": title,
                    "candidates": ordered,
                })
    return destination, themes


def _sort_distinctive_tokens(name):
    return {w for w in _sort_norm(name).split()
            if len(w) >= 4 and w not in SORT_GENERIC_WORDS}


def _sort_distinctive_join(name):
    return " ".join(sorted(_sort_distinctive_tokens(name)))


def _sort_match_tier(folder_name, candidate):
    """4 exact / space-insensitive, 3 word-boundary containment,
    2 fuzzy (whole names, or the distinctive words only -- catches
    misspellings like 'Buckingam palace', 'Kew gardins', 'Camben'),
    1 shared/fuzzy distinctive token, 0 no match. Thresholds tuned
    against a battery of real misspellings AND false-positive guards
    (East end vs West End, National gallery vs National Theatre must
    never match)."""
    f, c = _sort_norm(folder_name), _sort_norm(candidate)
    if not f or not c:
        return 0
    if f == c or f.replace(" ", "") == c.replace(" ", ""):
        return 4
    if re.search(r"\b" + re.escape(f) + r"\b", c) or \
       re.search(r"\b" + re.escape(c) + r"\b", f):
        return 3
    if SequenceMatcher(None, f, c).ratio() >= 0.88 or \
       SequenceMatcher(None, f.replace(" ", ""), c.replace(" ", "")).ratio() >= 0.88:
        return 2
    dfj, dcj = _sort_distinctive_join(folder_name), _sort_distinctive_join(candidate)
    if dfj and dcj and SequenceMatcher(None, dfj, dcj).ratio() >= 0.84:
        return 2
    toks_f = _sort_distinctive_tokens(folder_name)
    toks_c = _sort_distinctive_tokens(candidate)
    if toks_f & toks_c:
        return 1
    for ft in toks_f:
        for ct in toks_c:
            need = 0.80 if min(len(ft), len(ct)) >= 6 else 0.86
            if SequenceMatcher(None, ft, ct).ratio() >= need:
                return 1
    return 0


def _sort_best_match(folder_name, themes):
    if _sort_norm(folder_name) in SORT_FOLDER_STOPLIST:
        return 0, None, None
    best = (0, None, None)
    for theme in themes:
        for poi in theme["pois"]:
            for cand in poi["candidates"]:
                tier = _sort_match_tier(folder_name, cand)
                if tier > best[0]:
                    best = (tier, theme, poi)
    return best


def find_folder_by_name(folder, name):
    """Depth-first search of the bin tree for a folder by name
    (case-insensitive)."""
    if folder.GetName().strip().lower() == name.strip().lower():
        return folder
    for sub in folder.GetSubFolderList():
        found = find_folder_by_name(sub, name)
        if found:
            return found
    return None


def is_leaf_folder(folder):
    return not folder.GetSubFolderList()


def is_poi_folder(folder):
    """A POI folder is one whose subfolders are all leaves -- i.e. it
    directly holds the photographer/date bins (or nothing deeper).
    Containers like a city or region (R)/(C) folder have non-leaf
    children and are traversed through, never matched -- which is what
    stops a folder literally named 'London' being claimed by the
    'Tower of London' POI."""
    subs = folder.GetSubFolderList()
    return bool(subs) and all(is_leaf_folder(s) for s in subs)


def collect_poi_folders(container, out):
    """Depth-first: gather POI folders at any depth under a container,
    skipping already-created Theme/EXTRAS folders on reruns."""
    for child in container.GetSubFolderList():
        name = child.GetName()
        if name.startswith("Theme ") or name == SORT_EXTRAS_NAME:
            continue
        if is_poi_folder(child) or is_leaf_folder(child):
            out.append(child)
        else:
            collect_poi_folders(child, out)
    return out


def summary_dialog(title, lines):
    """Simple end-of-run report popup with an OK button."""
    dlg_disp = bmd.UIDispatcher(ui)
    dlg = dlg_disp.AddWindow(
        {
            "ID": "SummaryDlg",
            "WindowTitle": title,
            "Geometry": [240, 240, 420, 90 + 22 * len(lines)],
            "StyleSheet": PANEL_QSS,
        },
        [
            ui.VGroup(
                {"Spacing": 6},
                [ui.Label({"Text": line, "WordWrap": True}) for line in lines]
                + [ui.Button({"ID": "BtnSummaryOk", "Text": "OK", "Weight": 0})],
            )
        ],
    )

    def on_ok(_ev):
        dlg_disp.ExitLoop()

    dlg.On.BtnSummaryOk.Clicked = on_ok
    dlg.On.SummaryDlg.Close = on_ok
    hold_log_widget()
    dlg.Show()
    run_loop_resilient(dlg_disp, "summary dialog")
    dlg.Hide()
    release_log_widget()


def get_subfolder(parent, name):
    for sub in parent.GetSubFolderList():
        if sub.GetName() == name:
            return sub
    return None


def ensure_subfolder(media_pool, parent, name):
    existing = get_subfolder(parent, name)
    if existing:
        return existing
    return media_pool.AddSubFolder(parent, name)


def color_clips_recursive(folder, clip_color):
    """Set every clip's colour in a folder tree. Returns (ok, fail)."""
    ok = fail = 0
    for clip in folder.GetClipList():
        try:
            if clip.SetClipColor(clip_color):
                ok += 1
            else:
                fail += 1
        except Exception:
            fail += 1
    for sub in folder.GetSubFolderList():
        s_ok, s_fail = color_clips_recursive(sub, clip_color)
        ok += s_ok
        fail += s_fail
    return ok, fail


def try_set_folder_color(folder, color):
    """Bin label colours have no documented API -- attempt the plausible
    method names and report honestly. Harmless when unsupported."""
    for method_name in ("SetFolderColor", "SetClipColor", "SetColor"):
        try:
            method = getattr(folder, method_name)
            if method(color) is True:
                return True
        except Exception:
            continue
    return False


def sort_dialog(project, media_pool):
    """Blocking dialog: shoot-notes docx, destination folder, preview
    toggle. Returns dict or None."""
    dlg_disp = bmd.UIDispatcher(ui)
    result = {"params": None}

    root = media_pool.GetRootFolder()
    raw_folder = find_folder_by_name(root, SORT_PARENT_FOLDER)
    source_parent = raw_folder or root
    dest_names = [f.GetName() for f in source_parent.GetSubFolderList()
                  if not f.GetName().lower().startswith("theme")
                  and f.GetName() != SORT_EXTRAS_NAME]

    dlg = dlg_disp.AddWindow(
        {
            "ID": "SortDlg",
            "WindowTitle": "Sort by Shoot Notes",
            "Geometry": [200, 200, 480, 260],
            "StyleSheet": PANEL_QSS,
        },
        [
            ui.VGroup(
                {"Spacing": 8},
                [
                    ui.Label({"Text": "Shoot notes (.docx)"}),
                    ui.HGroup({"Spacing": 6, "Weight": 0}, [
                        ui.LineEdit({"ID": "SortNotesPath",
                                     "PlaceholderText": "Path to the shoot notes"}),
                        ui.Button({"ID": "BtnSortBrowse", "Text": "Browse", "Weight": 0}),
                    ]),
                    ui.Label({"Text": f"Destination folder (inside"
                                       f" {source_parent.GetName()})"}),
                    ui.ComboBox({"ID": "SortDestCombo"}),
                    ui.CheckBox({"ID": "SortPreview",
                                 "Text": "Preview only -- log the plan, move nothing",
                                 "Checked": True}),
                    ui.Label({"ID": "SortWarn", "Text": ""}),
                    ui.HGroup({"Spacing": 8, "Weight": 0}, [
                        ui.Button({"ID": "BtnSortCancel", "Text": "Cancel"}),
                        ui.Button({"ID": "BtnSortRun", "Text": "Run"}),
                    ]),
                ],
            )
        ],
    )

    ditems = dlg.GetItems()
    for name in dest_names:
        ditems["SortDestCombo"].AddItem(name)

    def on_browse(_ev):
        try:
            path = fusion.RequestFile()
            if path:
                ditems["SortNotesPath"].Text = str(path)
        except Exception:
            ditems["SortWarn"].Text = "File browser unavailable -- paste the path."

    def on_run(_ev):
        notes_path = (ditems["SortNotesPath"].Text or "").strip()
        if not notes_path or not os.path.isfile(notes_path):
            ditems["SortWarn"].Text = "Shoot notes file not found -- check the path."
            return
        if not notes_path.lower().endswith(".docx"):
            ditems["SortWarn"].Text = "Shoot notes must be .docx."
            return
        dest_name = ditems["SortDestCombo"].CurrentText
        if not dest_name:
            ditems["SortWarn"].Text = "No destination folder selected."
            return
        result["params"] = {
            "notes_path": notes_path,
            "dest_name": dest_name,
            "preview": bool(ditems["SortPreview"].Checked),
        }
        dlg_disp.ExitLoop()

    def on_cancel(_ev):
        dlg_disp.ExitLoop()

    dlg.On.BtnSortBrowse.Clicked = guard(on_browse)
    dlg.On.BtnSortRun.Clicked = guard(on_run)
    dlg.On.BtnSortCancel.Clicked = on_cancel
    dlg.On.SortDlg.Close = on_cancel

    hold_log_widget()
    dlg.Show()
    run_loop_resilient(dlg_disp, "setup dialog")
    dlg.Hide()
    release_log_widget()
    return result["params"]


def on_sort_shoot_notes(ev):
    project, _, media_pool = get_context()
    if not project:
        log("No project open.")
        return

    params = sort_dialog(project, media_pool)
    if not params:
        log("Sort by Shoot Notes cancelled.")
        return

    try:
        destination, themes = parse_shoot_notes(params["notes_path"])
    except ImportError:
        log("python-docx is not installed in Resolve's Python -- run the"
            " install_python_docx script.")
        return
    except Exception:
        log(f"Could not parse the shoot notes:\n{traceback.format_exc()}")
        return

    n_pois = sum(len(t["pois"]) for t in themes)
    log(f"Shoot notes parsed: destination '{destination}',"
        f" {len(themes)} theme(s), {n_pois} POI(s).")
    if not themes:
        log("No 'Theme N:' rows found -- is this a shoot-notes document?")
        return

    root = media_pool.GetRootFolder()
    raw_folder = find_folder_by_name(root, SORT_PARENT_FOLDER) or root
    dest_folder = get_subfolder(raw_folder, params["dest_name"])
    if not dest_folder:
        log(f"Destination folder '{params['dest_name']}' not found.")
        return

    # --- Find POI folders structurally, at any depth --------------------
    poi_folders = collect_poi_folders(dest_folder, [])
    log(f"Found {len(poi_folders)} POI folder(s) under"
        f" '{dest_folder.GetName()}'.")

    # --- Build the plan --------------------------------------------------
    plan = []   # (folder, theme or None, poi or None, tier)
    for sub in poi_folders:
        tier, theme, poi = _sort_best_match(sub.GetName(), themes)
        plan.append((sub, theme, poi, tier))

    tag = {4: "exact", 3: "contains", 2: "fuzzy", 1: "token"}
    matched_poi_keys = set()
    for theme in themes:
        color = SORT_FOLDER_COLORS[(theme["num"] - 1) % len(SORT_FOLDER_COLORS)]
        matched_here = [(f, p, t) for f, th, p, t in plan if th is theme]
        if not matched_here:
            continue
        log(f"Theme {theme['num']} - {theme['name']} [{color}]:")
        for folder, poi, tier in matched_here:
            matched_poi_keys.add((theme["num"], poi["label"]))
            log(f"    {folder.GetName()} -> {poi['label']} - "
                f"{poi['title'][:40]} ({tag[tier]})")
    extras_folders = [f for f, th, p, t in plan if th is None]
    if extras_folders:
        log(f"{SORT_EXTRAS_NAME}: "
            + ", ".join(f.GetName() for f in extras_folders))

    # POIs from the notes that no folder matched
    unfound_pois = []
    total_pois = 0
    for theme in themes:
        for poi in theme["pois"]:
            total_pois += 1
            if (theme["num"], poi["label"]) not in matched_poi_keys:
                unfound_pois.append(f"{poi['label']} - {poi['title'][:35]}")
    if unfound_pois:
        log(f"POIs in the notes with no matching footage ({len(unfound_pois)}):")
        for name in unfound_pois:
            log(f"    {name}")

    summary_lines = [
        f"Destination: {destination or params['dest_name']}",
        f"Themes in notes: {len(themes)}   POIs in notes: {total_pois}",
        f"POIs with footage found: {len(matched_poi_keys)}",
        f"Folders going to {SORT_EXTRAS_NAME}: {len(extras_folders)}",
    ]
    if unfound_pois:
        summary_lines.append("")
        summary_lines.append(f"Missed POIs / Checklist Items ({len(unfound_pois)}):")
        MAX_LISTED = 14  # keep the dialog a sane height on big shoots
        for name in unfound_pois[:MAX_LISTED]:
            summary_lines.append(f"  \u2022 {name} was missed.")
        if len(unfound_pois) > MAX_LISTED:
            summary_lines.append(f"  ...and {len(unfound_pois) - MAX_LISTED}"
                                 f" more (full list in the log).")
    else:
        summary_lines.append("Missed POIs / Checklist Items: none --"
                             " every POI in the notes has footage.")

    if params["preview"]:
        log("Preview only -- nothing was moved. Untick the preview box to"
            " apply this plan.")
        summary_dialog("Sort Preview", summary_lines
                       + ["", "PREVIEW ONLY -- nothing was moved."])
        return

    # --- Apply -----------------------------------------------------------
    moved = move_fails = 0
    folder_color_ok = folder_color_fail = 0
    clip_ok = clip_fail = 0
    theme_folders = {}

    for folder, theme, poi, tier in plan:
        if theme is None:
            continue  # unmatched POIs travel with their parent to EXTRAS
        theme_name = f"Theme {theme['num']} - {theme['name']}"
        theme_folder = theme_folders.get(theme_name)
        if not theme_folder:
            theme_folder = ensure_subfolder(media_pool, dest_folder, theme_name)
            theme_folders[theme_name] = theme_folder
        if not theme_folder:
            log(f"  Could not create '{theme_name}'.")
            move_fails += 1
            continue
        poi_title = poi["title"].split(" (")[0].strip()
        poi_name = f"{poi['label']} - {poi_title}"
        target_parent = ensure_subfolder(media_pool, theme_folder, poi_name)
        if not target_parent:
            move_fails += 1
            log(f"  Could not create target folder for {folder.GetName()}.")
            continue
        try:
            ok = media_pool.MoveFolders([folder], target_parent)
        except Exception:
            ok = False
        if ok:
            moved += 1
        else:
            move_fails += 1
            log(f"  Failed to move: {folder.GetName()}")

    # Everything left over (containers with only unmatched POIs, loose
    # unmatched folders) moves wholesale into EXTRAS, structure intact.
    extras_moved = 0
    extras_target = None
    for child in list(dest_folder.GetSubFolderList()):
        name = child.GetName()
        if name.startswith("Theme ") or name == SORT_EXTRAS_NAME:
            continue
        if extras_target is None:
            extras_target = ensure_subfolder(media_pool, dest_folder,
                                             SORT_EXTRAS_NAME)
        try:
            if extras_target and media_pool.MoveFolders([child], extras_target):
                extras_moved += 1
            else:
                log(f"  Failed to move to {SORT_EXTRAS_NAME}: {name}")
        except Exception:
            log(f"  Failed to move to {SORT_EXTRAS_NAME}: {name}")

    # --- Colour ------------------------------------------------------------
    for theme in themes:
        theme_name = f"Theme {theme['num']} - {theme['name']}"
        theme_folder = theme_folders.get(theme_name)
        if not theme_folder:
            continue
        color = SORT_FOLDER_COLORS[(theme["num"] - 1) % len(SORT_FOLDER_COLORS)]
        clip_color = SORT_CLIP_COLOR[color]

        def color_folders(folder):
            nonlocal folder_color_ok, folder_color_fail
            if try_set_folder_color(folder, color):
                folder_color_ok += 1
            else:
                folder_color_fail += 1
            for sub in folder.GetSubFolderList():
                color_folders(sub)

        color_folders(theme_folder)
        c_ok, c_fail = color_clips_recursive(theme_folder, clip_color)
        clip_ok += c_ok
        clip_fail += c_fail
        log(f"  {theme_name}: clips coloured {clip_color} ({c_ok} ok"
            + (f", {c_fail} failed" if c_fail else "") + ")")

    log(f"Sort done. Moved {moved} POI folder(s) ({move_fails} failed),"
        f" {extras_moved} folder(s) to {SORT_EXTRAS_NAME},"
        f" coloured {clip_ok} clip(s).")
    if folder_color_fail and not folder_color_ok:
        log("Note: bin label colours aren't scriptable on this build --"
            " right-click each Theme folder > colour to label them"
            " manually. All clips inside are already coloured.")

    summary_dialog("Sort Complete", summary_lines + [
        "",
        f"Moved {moved} POI folder(s), {extras_moved} to {SORT_EXTRAS_NAME}.",
        f"Coloured {clip_ok} clip(s).",
    ])


# ---------------------------------------------------------------------------
# Bin Finder -- search every bin by name/path and jump the Media Pool to
# it. Ported from the bin_finder_poc script: scans bin NAMES only (never
# GetClipList), so it stays fast on huge pools.
# ---------------------------------------------------------------------------
BIN_FINDER_RESULT_LIMIT = 100

# Pages that actually show a Media Pool. Jumping to a bin from anywhere
# else changes the pool's current folder with nothing on screen to show
# it -- which is exactly why "open from the Colour page" looked broken.
BIN_FINDER_POOL_PAGES = ("media", "cut", "edit")

# Where to land when a switch IS needed. Edit, not Media: this is an
# editor's tool, and the bin is wanted next to the timeline rather than in
# the ingest view.
BIN_FINDER_JUMP_PAGE = "edit"

# Arrow keys work in the RESULT LIST unconditionally (Tab into it from the
# search box). This flag additionally puts Up/Down/Enter on the search
# field itself, so you never have to leave it. It is off by default only
# because UIManager builds differ in whether enabling KeyPress on a text
# field still lets normal typing through -- if typing still works with this
# on, leave it on and you get arrows without tabbing.
BIN_FINDER_TRY_KEY_EVENTS = False

# Logs what triggered each open, and every key the list receives. Useful
# while confirming double-click/Enter actually arrive on a given build;
# turn off once it is all behaving.
BIN_FINDER_DEBUG_EVENTS = True

# Qt key codes, for the optional KeyPress path above.
_QT_KEY_UP = 16777235
_QT_KEY_DOWN = 16777237
_QT_KEY_RETURN = 16777220
_QT_KEY_ENTER = 16777221
_QT_KEY_ESCAPE = 16777216


def index_bins(media_pool, include_clips=False):
    """-> ordered list of entries for every bin -- and, optionally, every
    clip/file name too (so searches like '.png' find image files; opening
    a clip result opens its bin).

    Each entry carries the pre-lowercased name/path the scorer wants, the
    tree depth (used to break score ties in favour of shallower bins), and
    Resolve's own stale flag."""
    entries = []

    # The reference implementation refreshes first, and it matters in
    # collaboration mode: without it the walk can index folders Resolve
    # already considers gone.
    try:
        media_pool.RefreshFolders()
    except Exception:
        pass

    def walk(folder, parent_path="", depth=0):
        try:
            name = folder.GetName() or ("Media Pool" if depth == 0 else "Untitled Bin")
        except Exception:
            name = "Untitled Bin"
        path = f"{parent_path} / {name}" if parent_path else name
        try:
            is_stale = bool(folder.GetIsFolderStale())
        except Exception:
            is_stale = False
        entries.append({"name": name, "path": path, "folder": folder,
                        "kind": "bin", "depth": depth,
                        "name_lower": name.lower(), "path_lower": path.lower(),
                        "stale": is_stale})
        if include_clips:
            try:
                clips = folder.GetClipList() or []
            except Exception:
                clips = []
            for clip in clips:
                try:
                    clip_name = clip.GetName() or ""
                except Exception:
                    clip_name = ""
                if clip_name:
                    clip_path = f"{path} / {clip_name}"
                    entries.append({"name": clip_name, "path": clip_path,
                                    "folder": folder, "kind": "clip",
                                    "depth": depth + 1,
                                    "name_lower": clip_name.lower(),
                                    "path_lower": clip_path.lower(),
                                    "stale": is_stale})
        try:
            children = folder.GetSubFolderList() or []
        except Exception:
            children = []
        for child in children:
            walk(child, path, depth + 1)

    walk(media_pool.GetRootFolder())
    return entries


# --- Search scoring --------------------------------------------------------
# A direct port of the reference build's src/search.js, so both tools rank
# results identically. The behaviour worth keeping in mind: separators are
# flattened to spaces, so "ep3_audio" and "ep3 audio" are the same query;
# a multi-word query matches when EVERY word appears; and failing all
# that, a fuzzy subsequence match still finds "Buckingham" from "bkgm".
_BIN_SEP_RE = re.compile(r"[_\-./\\]+")
_BIN_WS_RE = re.compile(r"\s+")
_BIN_NON_ALNUM_RE = re.compile(r"[^a-z0-9]+")
_NEG_INF = float("-inf")


def _bin_normalize(value):
    text = str(value or "").strip().lower()
    text = _BIN_SEP_RE.sub(" ", text)
    return _BIN_WS_RE.sub(" ", text)


def _bin_compact(value):
    return _BIN_NON_ALNUM_RE.sub("", _bin_normalize(value))


def _bin_tokenize(value):
    return [token for token in _bin_normalize(value).split(" ") if token]


def _bin_fuzzy_score(needle, haystack):
    """Subsequence match: every character of `needle` must appear in
    `haystack` in order. Scores tighter runs higher. -inf when it does not
    match at all."""
    if not needle:
        return 0
    if not haystack:
        return _NEG_INF
    last_index = -1
    gaps = 0
    for char in needle:
        index = haystack.find(char, last_index + 1)
        if index == -1:
            return _NEG_INF
        if last_index != -1:
            gaps += index - last_index - 1
        last_index = index
    return max(0, 1200 - gaps * 8 - (len(haystack) - len(needle)))


def _bin_score_entry(entry, query):
    """Higher is better; -inf means no match at all."""
    if not query:
        return 1

    name = entry.get("name_lower") or _bin_normalize(entry["name"])
    path = entry.get("path_lower") or _bin_normalize(entry["path"])
    query_compact = _bin_compact(query)
    name_compact = _bin_compact(name)
    path_compact = _bin_compact(path)
    tokens = _bin_tokenize(query)

    score = _NEG_INF

    if name == query:
        score = max(score, 10000)
    if name.startswith(query):
        score = max(score, 9000 - len(name))
    if query in name:
        score = max(score, 8000 - name.find(query) * 10)
    if query in path:
        score = max(score, 6500 - path.find(query) * 4)

    if len(tokens) > 1:
        if all(token in name for token in tokens):
            score = max(score, 7600 - len(name))
        if all(token in path for token in tokens):
            score = max(score, 6200 - len(path))

    name_fuzzy = _bin_fuzzy_score(query_compact, name_compact)
    if name_fuzzy > _NEG_INF:
        score = max(score, 4200 + name_fuzzy)

    path_fuzzy = _bin_fuzzy_score(query_compact, path_compact)
    if path_fuzzy > _NEG_INF:
        score = max(score, 3000 + path_fuzzy)

    return score


def group_matches_by_hierarchy(matches):
    """Re-order ranked matches so a matching bin is followed immediately by
    its OWN matching subfolders, instead of every subfolder clumping
    together further down the list.

    Searching "temple" should read:

        TEMPLE A
          - TEMPLE A / Stalls
        TEMPLE B

    not TEMPLE A, TEMPLE B, then both sets of subfolders where you cannot
    tell which parent they belong to. The reference build shows a flat
    ranked list; this is deliberately the older Infinite Forms behaviour,
    because a location-per-bin pipeline reads far better grouped.

    Best-scoring bins still come first -- ranking decides the order of the
    GROUPS, and each group carries its own children along."""
    by_path = {entry["path"]: entry for entry in matches}

    def has_matching_ancestor(entry):
        parts = entry["path"].split(" / ")
        for cut in range(1, len(parts)):
            if " / ".join(parts[:cut]) in by_path:
                return True
        return False

    ordered = []
    emitted = set()

    def emit(entry, indent):
        if entry["path"] in emitted:
            return
        emitted.add(entry["path"])
        row = dict(entry)
        row["indent"] = indent
        ordered.append(row)
        prefix = entry["path"] + " / "
        children = [candidate for candidate in matches
                    if candidate["path"].startswith(prefix)]
        # Shallowest first, then alphabetical: a direct child pulls its own
        # subtree along before the next sibling starts.
        children.sort(key=lambda c: (c.get("depth", 0), c["path"]))
        for child in children:
            emit(child, indent + 1)

    for entry in matches:              # ranked order drives the groups
        if entry["path"] in emitted or has_matching_ancestor(entry):
            continue
        emit(entry, 0)

    # Belt and braces -- never silently drop a result.
    for entry in matches:
        if entry["path"] not in emitted:
            emit(entry, 0)

    return ordered


def search_bins(bins, query, limit=BIN_FINDER_RESULT_LIMIT):
    """Ranked matches, best first, then regrouped so subfolders sit under
    the parent they belong to. Ties go to the shallower bin, then to
    alphabetical path -- so a top-level bin beats one buried six levels
    down when they score the same."""
    normalized = _bin_normalize(query)
    scored = []
    for entry in bins:
        score = _bin_score_entry(entry, normalized)
        if score > _NEG_INF:
            scored.append((score, entry))
    scored.sort(key=lambda pair: (-pair[0], pair[1].get("depth", 0),
                                 pair[1]["path"]))
    return group_matches_by_hierarchy([entry for _score, entry
                                       in scored[:limit]])


BIN_FINDER = {"win": None, "items": None, "bins": [], "by_path": {},
              "matches": [], "tree_items": [], "selected": 0,
              "project_id": None, "indexed_at": None,
              "pinned": False, "compact": False, "shown_results": None}

BIN_FINDER_FULL_SIZE = [560, 460]
# Compact is a single row -- search box plus Expand/Pin -- and stays that
# small until you type something, at which point it grows to show the
# results and shrinks again when the query is cleared. Same behaviour as
# the reference build's 58px -> 320px window.
BIN_FINDER_COMPACT_SIZE = [560, 96]
BIN_FINDER_COMPACT_QUERY_SIZE = [560, 380]


def _bin_finder_project_id(project):
    try:
        return project.GetUniqueId()
    except Exception:
        try:
            return project.GetName()
        except Exception:
            return None


def _bin_finder_status(text):
    try:
        BIN_FINDER["items"]["BinFinderStatus"].Text = text
    except Exception:
        pass


def _bin_finder_apply_selection():
    """Push BIN_FINDER["selected"] into the tree widget so the highlight
    follows the Up/Down buttons, and scroll it into view."""
    items_list = BIN_FINDER["tree_items"]
    if not items_list:
        return
    index = max(0, min(BIN_FINDER["selected"], len(items_list) - 1))
    BIN_FINDER["selected"] = index
    item = items_list[index]
    tree = BIN_FINDER["items"]["BinTree"]
    for setter in (
        lambda: tree.SetCurrentItem(item),
        lambda: setattr(item, "Selected", True),
    ):
        try:
            setter()
        except Exception:
            pass
    try:
        tree.ScrollToItem(item)
    except Exception:
        pass


def _bin_finder_current_entry():
    """The entry to act on: whatever the tree says is current (so a mouse
    click wins), falling back to our own tracked index."""
    matches = BIN_FINDER["matches"]
    if not matches:
        return None
    try:
        item = BIN_FINDER["items"]["BinTree"].CurrentItem()
        path = item.Text[2] if item else None
        if path:
            for index, entry in enumerate(matches):
                if entry["path"] == path:
                    BIN_FINDER["selected"] = index
                    return entry
    except Exception:
        pass
    index = max(0, min(BIN_FINDER["selected"], len(matches) - 1))
    return matches[index]


def _bin_finder_move(delta):
    """Move the highlight. Anchored on the current entry first, so arrows
    continue from wherever the user last clicked."""
    if not BIN_FINDER["matches"]:
        return
    _bin_finder_current_entry()      # syncs "selected" to the tree
    BIN_FINDER["selected"] += delta
    _bin_finder_apply_selection()


def _bin_finder_wants_results():
    """Compact mode shows nothing until you type -- that is the whole
    point of it. Full mode always shows the list."""
    if not BIN_FINDER["compact"]:
        return True
    try:
        return bool((BIN_FINDER["items"]["BinQuery"].Text or "").strip())
    except Exception:
        return True


def _bin_finder_target_size():
    if not BIN_FINDER["compact"]:
        return BIN_FINDER_FULL_SIZE
    return (BIN_FINDER_COMPACT_QUERY_SIZE if _bin_finder_wants_results()
            else BIN_FINDER_COMPACT_SIZE)


def _bin_finder_resize(size):
    """Grow/shrink the open window. Resize() alone was reported as a no-op
    on this Resolve build, so the Geometry property is set as well -- one
    of the two lands on any given build, and both are harmless."""
    window = BIN_FINDER["win"]
    if not window:
        return
    try:
        window.Resize(size)
    except Exception:
        pass
    try:
        geo = window.Geometry
        window.Geometry = [geo[0], geo[1], size[0], size[1]]
    except Exception:
        pass
    try:
        window.RecalcLayout()
    except Exception:
        pass


def _bin_finder_refresh(_ev=None):
    """Re-run the search and repaint the result list. Wired to the query
    field's TextChanged, so results follow every keystroke the way the
    reference build does."""
    ditems = BIN_FINDER["items"]
    tree = ditems["BinTree"]
    query = (ditems["BinQuery"].Text or "")
    # Compact with an empty query has nothing to show, so do not even
    # search -- same as the reference build emptying its filtered list.
    matches = (search_bins(BIN_FINDER["bins"], query)
               if _bin_finder_wants_results() else [])
    BIN_FINDER["matches"] = matches
    BIN_FINDER["by_path"] = {}
    BIN_FINDER["tree_items"] = []
    try:
        tree.Clear()
    except Exception:
        pass
    for entry in matches:
        BIN_FINDER["by_path"][entry["path"]] = entry["folder"]
        if entry.get("kind") == "clip":
            ext = os.path.splitext(entry["name"])[1]
            type_label = ext.upper() if ext else "File"
        else:
            type_label = "Bin"
        if entry.get("stale"):
            type_label += " (stale)"
        # Indent children under the parent they matched beneath, so the
        # grouping is visible and not just an ordering you have to infer.
        indent = entry.get("indent", 0)
        label = (("    " * (indent - 1) + "└ " + entry["name"])
                 if indent else entry["name"])
        try:
            item = tree.NewItem()
            item.Text[0] = label
            item.Text[1] = type_label
            item.Text[2] = entry["path"]
            tree.AddTopLevelItem(item)
            BIN_FINDER["tree_items"].append(item)
        except Exception:
            pass

    # Always leave something selected, so Enter works the instant you stop
    # typing without having to click a row first.
    BIN_FINDER["selected"] = 0
    _bin_finder_apply_selection()

    # In compact mode, crossing between "empty query" and "has a query"
    # is what grows and shrinks the window. Only act on the transition --
    # resizing on every keystroke would fight the person typing.
    if _bin_finder_wants_results() != BIN_FINDER.get("shown_results"):
        _bin_finder_apply_compact()

    n_bins = sum(1 for e in BIN_FINDER["bins"] if e.get("kind") == "bin")
    n_clips = sum(1 for e in BIN_FINDER["bins"] if e.get("kind") == "clip")
    scanned = BIN_FINDER["indexed_at"]
    _bin_finder_status(
        f"{n_bins} bin(s)"
        + (f" + {n_clips} file(s)" if n_clips else "")
        + f" indexed -- {len(matches)} match(es)"
        + (f" (capped at {BIN_FINDER_RESULT_LIMIT})"
           if len(matches) == BIN_FINDER_RESULT_LIMIT else "")
        + (f" -- last scan {scanned}" if scanned else "")
    )


def _bin_finder_reindex(_ev=None):
    project, _, media_pool = get_context()
    if not media_pool:
        _bin_finder_status("No project open.")
        return
    include_clips = False
    try:
        include_clips = bool(BIN_FINDER["items"]["BinIncludeClips"].Checked)
    except Exception:
        pass
    BIN_FINDER["bins"] = index_bins(media_pool, include_clips=include_clips)
    BIN_FINDER["project_id"] = _bin_finder_project_id(project)
    BIN_FINDER["indexed_at"] = time.strftime("%H:%M", time.localtime())
    _bin_finder_refresh()


def _bin_finder_open(_ev=None):
    """Jump the Media Pool to the selected bin.

    Switches to the Edit page first when the current page has no Media Pool
    on it. Without that, opening a bin from Colour or Fairlight changed the
    pool's current folder with nothing on screen to prove it -- which is
    indistinguishable from the button doing nothing at all.

    Already on Media, Cut or Edit? Then nothing is switched -- the pool is
    right there, and yanking someone off the page they chose would be
    worse than leaving them on it."""
    project, _, media_pool = get_context()
    if not project or not media_pool:
        _bin_finder_status("No project open.")
        return

    # An index built against a different project holds folder references
    # that no longer mean anything.
    current_id = _bin_finder_project_id(project)
    if BIN_FINDER["project_id"] and current_id != BIN_FINDER["project_id"]:
        _bin_finder_status("Project changed -- re-scanning, then try again.")
        log("Bin Finder: project changed since the last scan, re-indexing.")
        _bin_finder_reindex()
        return

    entry = _bin_finder_current_entry()
    if not entry:
        _bin_finder_status("Select a result first.")
        return

    if entry.get("stale"):
        _bin_finder_status("That bin is stale in Resolve -- click Re-scan.")
        return

    folder = entry["folder"]
    path = entry["path"]

    # --- Make sure the Media Pool is visible ------------------------------
    switched_from = None
    try:
        page = (resolve.GetCurrentPage() or "").lower()
    except Exception:
        page = ""
    if page and page not in BIN_FINDER_POOL_PAGES:
        try:
            if resolve.OpenPage(BIN_FINDER_JUMP_PAGE):
                switched_from = page
        except Exception:
            pass

    try:
        ok = media_pool.SetCurrentFolder(folder)
    except Exception:
        ok = False

    if ok:
        jumped = (f" (switched from {switched_from} to"
                  f" {BIN_FINDER_JUMP_PAGE})" if switched_from else "")
        _bin_finder_status(f"Opened: {path}{jumped}")
        log(f"Bin Finder: opened {path}"
            + (f" -- switched from the {switched_from} page to"
               f" {BIN_FINDER_JUMP_PAGE}" if switched_from else ""))
        # Compact is a launcher: once it has done its job, clear the query
        # so the window collapses back to the single search row.
        if BIN_FINDER["compact"]:
            _bin_finder_clear()
    else:
        _bin_finder_status("Resolve failed to switch bins -- click Re-scan.")


def _bin_finder_clear(_ev=None):
    try:
        BIN_FINDER["items"]["BinQuery"].Text = ""
    except Exception:
        pass
    _bin_finder_refresh()


def _bin_finder_key(ev=None):
    """Qt KeyPress handling. Always active on the result list, where it
    cannot interfere with anything; optional on the search field, where
    UIManager builds differ in whether enabling the event still lets
    normal typing through (BIN_FINDER_TRY_KEY_EVENTS)."""
    key = None
    for reader in (lambda: ev["Key"], lambda: ev.get("Key"),
                   lambda: getattr(ev, "Key", None)):
        try:
            key = reader()
            if key is not None:
                break
        except Exception:
            continue
    if BIN_FINDER_DEBUG_EVENTS:
        log(f"Bin Finder: key event, Key={key!r}")
    if key == _QT_KEY_DOWN:
        _bin_finder_move(1)
    elif key == _QT_KEY_UP:
        _bin_finder_move(-1)
    elif key in (_QT_KEY_RETURN, _QT_KEY_ENTER):
        _bin_finder_open()
    elif key == _QT_KEY_ESCAPE:
        _bin_finder_clear()


def _bin_finder_open_via(source):
    """Wrap the open action with a note about what triggered it. Cheap, and
    it turns "double-click does nothing" into a one-line answer: either the
    Console shows the trigger fired and the open failed, or it shows
    nothing and the event never arrived."""
    def handler(_ev=None):
        if BIN_FINDER_DEBUG_EVENTS:
            log(f"Bin Finder: open requested via {source}.")
        _bin_finder_open()
    handler.__name__ = f"_bin_finder_open_{source.replace(' ', '_')}"
    return handler


def _bin_finder_hide(_ev=None):
    try:
        BIN_FINDER["win"].Hide()
    except Exception:
        pass


def _bin_finder_apply_compact():
    """Compact = one row: the search box plus Expand and Pin. The title,
    the Find/Re-scan buttons, the checkbox and the footer all go away, and
    the result list only appears once there is something to show -- which
    is also when the window grows.

    The search box itself is never hidden. It living inside the old
    collapsible body is exactly what made compact mode useless before."""
    compact = BIN_FINDER["compact"]
    show_results = _bin_finder_wants_results()
    for widget_id, hidden in (
        ("BinHeader", compact),
        ("BtnBinSearch", compact),
        ("BtnBinReindex", compact),
        ("BinIncludeClips", compact),
        ("BinActions", compact),
        ("BinResults", not show_results),
    ):
        try:
            BIN_FINDER["items"][widget_id].Hidden = hidden
        except Exception:
            pass
    try:
        BIN_FINDER["items"]["BtnBinCompact"].Text = ("Expand" if compact
                                                     else "Compact")
    except Exception:
        pass
    BIN_FINDER["shown_results"] = show_results
    _bin_finder_resize(_bin_finder_target_size())


def _bin_finder_toggle_compact(_ev=None):
    """Rebuild with the new geometry baked in -- creation geometry is the
    one sizing mechanism proven to work on this build, and Pin already
    relies on it.

    Entering compact clears the query, so the window opens at its small
    single-row size rather than immediately expanding. Leaving compact
    keeps whatever was typed."""
    BIN_FINDER["compact"] = not BIN_FINDER["compact"]
    query = ""
    if not BIN_FINDER["compact"]:
        try:
            query = BIN_FINDER["items"]["BinQuery"].Text or ""
        except Exception:
            pass
    BIN_FINDER["shown_results"] = None
    _bin_finder_hide()
    BIN_FINDER["win"] = None
    _bin_finder_build(query)


def _bin_finder_toggle_pin(_ev=None):
    """Always-on-top applies at creation, so pinning rebuilds the finder
    window (preserving the query text)."""
    BIN_FINDER["pinned"] = not BIN_FINDER["pinned"]
    query = ""
    try:
        query = BIN_FINDER["items"]["BinQuery"].Text or ""
    except Exception:
        pass
    _bin_finder_hide()
    BIN_FINDER["win"] = None
    _bin_finder_build(query)
    log("Bin Finder pinned on top." if BIN_FINDER["pinned"]
        else "Bin Finder unpinned.")


def _bin_finder_build(initial_query=""):
    props = {
        "ID": "BinFinderWin",
        "WindowTitle": "Bin Finder",
        "Geometry": [220, 220] + (BIN_FINDER_COMPACT_SIZE
                                  if BIN_FINDER["compact"]
                                  else BIN_FINDER_FULL_SIZE),
        "StyleSheet": PANEL_QSS,
    }
    if BIN_FINDER["pinned"]:
        props["WindowFlags"] = {"Window": True, "WindowStaysOnTopHint": True}

    query_props = {"ID": "BinQuery", "Text": initial_query,
                   "PlaceholderText": "Search bins (and files)..."}
    if BIN_FINDER_TRY_KEY_EVENTS:
        query_props["Events"] = {"KeyPress": True}

    finder = disp.AddWindow(
        props,
        [
            ui.VGroup(
                {"Spacing": 8},
                [
                    ui.HGroup({"ID": "BinHeader", "Spacing": 6, "Weight": 0}, [
                        ui.Label({"Text": "<b>Bin Finder</b>", "Weight": 1}),
                    ]),
                    # Search row: OUTSIDE anything collapsible (the search
                    # box being inside the old collapsible body is what
                    # broke compact mode), and it carries Expand + Pin so
                    # compact really is a single row.
                    ui.HGroup({"Spacing": 6, "Weight": 0}, [
                        ui.LineEdit(query_props),
                        # Kept for Full mode: if this Resolve build does
                        # not deliver TextChanged, this button is the only
                        # way left to search.
                        ui.Button({"ID": "BtnBinSearch", "Text": "Find",
                                   "Weight": 0}),
                        ui.Button({"ID": "BtnBinReindex", "Text": "Re-scan",
                                   "Weight": 0}),
                        ui.Button({"ID": "BtnBinCompact",
                                   "Text": "Expand" if BIN_FINDER["compact"]
                                           else "Compact",
                                   "Weight": 0}),
                        ui.Button({"ID": "BtnBinPin",
                                   "Text": "Unpin" if BIN_FINDER["pinned"] else "Pin",
                                   "Weight": 0}),
                    ]),
                    ui.CheckBox({"ID": "BinIncludeClips",
                                 "Text": "Include file names (slower scan --"
                                         " search '.png', '.wav' etc.)",
                                 "Checked": False, "Weight": 0}),
                    ui.VGroup({"ID": "BinResults", "Spacing": 6, "Weight": 1}, [
                        # Events must be ENABLED here, not merely handled.
                        # Registering On.BinTree.ItemDoubleClicked without
                        # this is why double-click did nothing: the signal
                        # was never delivered. KeyPress is safe to enable on
                        # a Tree (unlike on the search field) because there
                        # is no typing here to interfere with -- that is what
                        # gives Up/Down/Enter once the list has focus.
                        ui.Tree({"ID": "BinTree", "SortingEnabled": False,
                                 "AlternatingRowColors": True, "Weight": 1,
                                 "Events": {"ItemDoubleClicked": True,
                                            "ItemActivated": True,
                                            "ItemClicked": True,
                                            "CurrentItemChanged": True,
                                            "KeyPress": True}}),
                        ui.Label({"ID": "BinFinderStatus", "Text": "",
                                  "Weight": 0}),
                    ]),
                    ui.HGroup({"ID": "BinActions", "Spacing": 8, "Weight": 0}, [
                        ui.Button({"ID": "BtnBinClose", "Text": "Close"}),
                        ui.Button({"ID": "BtnBinOpen", "Text": "Open Bin"}),
                    ]),
                ],
            )
        ],
    )
    BIN_FINDER["win"] = finder
    BIN_FINDER["items"] = finder.GetItems()

    tree = BIN_FINDER["items"]["BinTree"]
    try:
        tree.ColumnCount = 3
        header = tree.NewItem()
        header.Text[0] = "Name"
        header.Text[1] = "Type"
        header.Text[2] = "Path"
        tree.SetHeaderItem(header)
        tree.ColumnWidth[0] = 160
        tree.ColumnWidth[1] = 60
    except Exception:
        pass

    # Preemptive no-op handlers for the chatty events we do NOT use, so
    # nothing reaches the dispatcher unregistered (KeyError: 'On').
    def _noop(_ev=None):
        pass
    for event_name in ("TextEdited", "EditingFinished"):
        try:
            setattr(finder.On.BinQuery, event_name, _noop)
        except Exception:
            pass
    for event_name in ("ItemSelectionChanged", "ItemChanged"):
        try:
            setattr(finder.On.BinTree, event_name, _noop)
        except Exception:
            pass

    # Clicking a row syncs our tracked selection, so Enter and the Open Bin
    # button always act on the row you can see highlighted.
    def _bin_finder_sync(_ev=None):
        _bin_finder_current_entry()
    for event_name in ("ItemClicked", "CurrentItemChanged"):
        try:
            setattr(finder.On.BinTree, event_name, guard(_bin_finder_sync))
        except Exception:
            pass

    finder.On.BtnBinPin.Clicked = guard(_bin_finder_toggle_pin)
    finder.On.BtnBinCompact.Clicked = guard(_bin_finder_toggle_compact)
    finder.On.BtnBinSearch.Clicked = guard(_bin_finder_refresh)
    finder.On.BtnBinReindex.Clicked = guard(_bin_finder_reindex)
    finder.On.BtnBinOpen.Clicked = guard(_bin_finder_open_via("Open Bin button"))
    finder.On.BtnBinClose.Clicked = guard(_bin_finder_hide)
    finder.On.BinFinderWin.Close = guard(_bin_finder_hide)

    # Toggling file-name indexing re-scans on the spot, rather than
    # silently doing nothing until Re-scan is clicked. Registering it also
    # means its Toggled event is never unhandled.
    try:
        finder.On.BinIncludeClips.Toggled = guard(_bin_finder_reindex)
    except Exception:
        pass

    # Search as you type, like the reference build -- the Find button is
    # now a fallback rather than the only way in.
    try:
        finder.On.BinQuery.TextChanged = guard(_bin_finder_refresh)
    except Exception:
        pass
    # Enter in the search box OPENS the highlighted result. It used to
    # re-run the search, which is why pressing Enter looked like it did
    # nothing -- the search had already run on the previous keystroke.
    try:
        finder.On.BinQuery.ReturnPressed = guard(
            _bin_finder_open_via("Enter in search box"))
    except Exception:
        pass
    # Double-click, plus ItemActivated (which Qt also fires on Enter).
    # These only arrive because the Tree enables them in its Events dict.
    for event_name, label in (("ItemDoubleClicked", "list double-click"),
                              ("ItemActivated", "list activate")):
        try:
            setattr(finder.On.BinTree, event_name,
                    guard(_bin_finder_open_via(label)))
        except Exception:
            pass
    # Up/Down/Enter/Esc once the list has focus -- Tab out of the search box
    # to get there. Safe to enable unconditionally on a Tree.
    try:
        finder.On.BinTree.KeyPress = guard(_bin_finder_key)
    except Exception:
        pass
    if BIN_FINDER_TRY_KEY_EVENTS:
        try:
            finder.On.BinQuery.KeyPress = guard(_bin_finder_key)
        except Exception:
            pass

    _bin_finder_reindex()
    _bin_finder_apply_compact()
    finder.Show()
    try:
        BIN_FINDER["items"]["BinQuery"].SetFocus()
    except Exception:
        pass


def on_bin_finder(ev):
    """Modeless, on the MAIN dispatcher (a nested dispatcher crashes on
    main-window events). Reopening re-indexes and reuses the window."""
    project, _, media_pool = get_context()
    if not project:
        log("No project open.")
        return
    if BIN_FINDER["win"] is not None:
        _bin_finder_reindex()
        BIN_FINDER["win"].Show()
        return
    _bin_finder_build()


# ---------------------------------------------------------------------------
# Auto-updater -- checks the GitHub repo's VERSION file at startup and,
# when it differs from BUILD_TAG, offers a one-click self-update. The
# downloaded file is VALIDATED before anything is touched (size sanity +
# full compile -- a truncated download can't brick the install), the old
# file is backed up alongside, and the swap is atomic.
# ---------------------------------------------------------------------------
UPDATE_STATE = {"available": False, "remote": "", "startup": None}


def _update_raw_url(filename):
    return (f"https://raw.githubusercontent.com/{UPDATE_REPO}/"
            f"{UPDATE_BRANCH}/{filename}")


def _update_repo_page_url():
    """The human-facing repo page -- what we tell people to go download
    from, as opposed to the raw file URLs used for fetching."""
    return f"https://github.com/{UPDATE_REPO}"


def _fetch_remote_version(timeout):
    """Read the repo's VERSION file. Returns (status, version):

        ("ok", "2026-07-31.1")  -- read it
        ("private", None)       -- HTTP 404: the repo is private, renamed,
                                   or UPDATE_REPO is wrong. GitHub returns
                                   404 rather than 403 for private repos
                                   when you are not authenticated, so this
                                   is the shape a private repo takes.
        ("certs", None)         -- TLS worked but the certificate could not
                                   be verified. On macOS this almost always
                                   means a python.org Python whose CA
                                   bundle was never installed (its
                                   "Install Certificates.command" was never
                                   run). Everything else on the machine can
                                   reach the internet fine, which makes
                                   this very easy to misread as offline.
        ("offline", None)       -- no network, DNS, timeout, anything else

    Telling these apart matters, because each has a completely different
    fix and only one of them is "check your wifi". Never raises -- an
    update check must not be able to break anything."""
    if not UPDATE_REPO:
        return ("offline", None)
    try:
        with urllib.request.urlopen(_update_raw_url("VERSION"),
                                    timeout=timeout) as response:
            return ("ok", response.read().decode("utf-8", "replace").strip())
    except urllib.error.HTTPError as err:
        # HTTPError subclasses URLError, so it has to be caught first.
        return ("private" if err.code == 404 else "offline", None)
    except urllib.error.URLError as err:
        if isinstance(getattr(err, "reason", None), ssl.SSLError):
            return ("certs", None)
        return ("offline", None)
    except ssl.SSLError:
        return ("certs", None)
    except Exception:
        return ("offline", None)


def _update_result(status, remote):
    """One place that turns a check result into (log line, dialog title,
    dialog lines) -- so the button and the launch check always say exactly
    the same thing.

    PLAIN TEXT ONLY. summary_dialog puts each line in a ui.Label, and Qt
    decides plain-vs-rich per label by looking for a tag: a line carrying
    an HTML entity but no tag renders the entity literally, which is where
    the stray "&nbsp;" came from. Keep lines under ~52 characters so they
    fit the dialog without wrapping."""
    repo_url = _update_repo_page_url()

    if status == "unconfigured":
        return ("Update checking is not configured (UPDATE_REPO is empty).",
                "Check for Update",
                ["Update checking is not set up.",
                 "",
                 "UPDATE_REPO is empty near the top of the plugin",
                 "file, so there is no repo to check against."])

    if status == "private":
        return ("Repo set to private -- update checking will not work until"
                " it is made public.",
                "Repo Set To Private",
                ["Repo set to private.",
                 "",
                 repo_url,
                 "",
                 "GitHub returned 404, which is what a private repo",
                 "looks like to anyone not signed in to it.",
                 "Update checking needs the repo to be public.",
                 "",
                 "Nothing on this machine can fix it -- the repo",
                 "owner has to make it public.",
                 "",
                 "(A renamed repo, or a wrong UPDATE_REPO in the",
                 "plugin file, looks identical from out here.)"])

    if status == "certs":
        return ("HTTPS certificates are not installed for this Python, so the"
                " update check could not run.",
                "Certificates Missing",
                ["Cannot verify GitHub's certificate.",
                 "",
                 "This is not a network problem -- the rest of the",
                 "machine can reach the internet fine. The Python",
                 "Resolve uses was installed from python.org",
                 "without its root certificates.",
                 "",
                 "Fix it once:",
                 "Open the Applications folder, open the",
                 "\"Python 3.x\" folder inside it, and double-click",
                 "\"Install Certificates.command\".",
                 "",
                 "Then check again."])

    if status != "ok":
        return ("Could not reach GitHub -- this machine appears to be offline.",
                "Check for Update",
                ["Could not reach GitHub.",
                 "",
                 "Tried:",
                 repo_url,
                 "",
                 "No answer at all, so this machine is probably",
                 "offline. Check your connection and try again.",
                 "Nothing was changed."])

    if not remote:
        return ("The repo's VERSION file is empty -- nothing to compare to.",
                "Check for Update",
                ["Could not read a version from the repo.",
                 "",
                 "The VERSION file at the repo root is empty.",
                 "Whoever published the last release needs to",
                 "fill it in."])

    if remote == BUILD_TAG:
        return (f"Up to date -- {BUILD_TAG} is the latest build.",
                "Up To Date",
                ["You are up to date.",
                 "",
                 f"Installed build:   {BUILD_TAG}",
                 f"Latest release:    {remote}"])

    lines = ["A newer build is available.",
             "",
             f"You have:          {BUILD_TAG}",
             f"Latest release:    {remote}",
             "",
             "To update, download the project from:",
             repo_url,
             "",
             "then double-click",
             "\"One Click Install Infinite Forms.command\".",
             "",
             "It backs up the version you have now before",
             "replacing it."]
    if UPDATE_STATE["available"]:
        lines += ["",
                  "Or use the Update button in the panel header",
                  "to install it in place."]
    return (f"Update available: {remote} (this is {BUILD_TAG}).",
            "Update Available", lines)


def check_for_update():
    """Automatic check, run once at launch. Never blocks long, never raises.
    The 3-second timeout is deliberate: a slow network must not delay the
    panel opening.

    Runs before the window exists, which is fine -- log() buffers into
    LOG_LINES and build_main_panel() flushes the whole trail into the log
    box the moment it is created. The result is stashed so the panel can
    pop the same dialog the button would, once there is a window to put it
    in front of (see show_startup_update_notice)."""
    UPDATE_STATE["startup"] = None
    if not UPDATE_REPO:
        return

    status, remote = _fetch_remote_version(3)
    message, title, lines = _update_result(status, remote)

    if status == "ok" and remote and remote != BUILD_TAG:
        UPDATE_STATE["available"] = True
        UPDATE_STATE["remote"] = remote

    if status == "ok" and remote == BUILD_TAG:
        # Confirm it in the panel log. No dialog -- nobody needs a popup to
        # say nothing changed -- but silence here is indistinguishable from
        # "the check never ran", which is exactly how this read before.
        log(f"Update check: up to date ({BUILD_TAG}).")
        return
    if status == "offline":
        # Transient and not actionable -- Console only, so a machine that
        # is offline on purpose is not nagged at every launch.
        print("[Infinite Forms] update check skipped (offline).")
        return

    log(f"Update check: {message}")
    UPDATE_STATE["startup"] = (title, lines)


def show_startup_update_notice():
    """Show the launch check's result, once the main panel exists so the
    dialog has something to sit in front of. Set UPDATE_STARTUP_DIALOG to
    False to keep launch findings in the log only."""
    if not UPDATE_STARTUP_DIALOG:
        return
    notice = UPDATE_STATE.get("startup")
    if not notice:
        return
    UPDATE_STATE["startup"] = None
    title, lines = notice
    try:
        summary_dialog(title, lines)
    except Exception:
        log(f"Could not show the update dialog:\n{traceback.format_exc()}")


def on_check_for_update(_ev=None):
    """Manual update check, driven by the panel's Check for Update button.

    Notify-only by design: it reads the repo's VERSION file, reports what
    it finds, and tells you where to download from. It never writes to the
    installed plugin -- that is the header Update button's job, and that
    one only appears when the launch check already found something.

    Deliberately does NOT set UPDATE_STATE either: the header button is
    built once, when the window is built, so flipping that flag here would
    promise a button that cannot appear until the next launch."""
    if not UPDATE_REPO:
        message, title, lines = _update_result("unconfigured", None)
        log(message)
        summary_dialog(title, lines)
        return

    log(f"Checking {UPDATE_REPO} for a newer build...")
    # A manual click can afford a longer wait than the launch check -- the
    # person is sitting there having asked for it. The panel does freeze for
    # the duration; that is the price of UIManager being single-threaded,
    # and 8s is the worst case.
    status, remote = _fetch_remote_version(8)
    message, title, lines = _update_result(status, remote)
    log(message)
    summary_dialog(title, lines)


# ---------------------------------------------------------------------------
# Wire it up and run
# ---------------------------------------------------------------------------
def on_close(ev):
    disp.ExitLoop()


try:
    check_for_update()
    build_main_panel()
    print("[Infinite Forms] showing panel")
    # After the panel exists, so the dialog has a window to sit in front of
    # and its nested dispatcher is not the only thing on screen.
    show_startup_update_notice()
    run_loop_resilient(disp, "main panel")
    try:
        win.Hide()
    except Exception:
        pass
    print("[Infinite Forms] panel closed normally")
except Exception:
    print("[Infinite Forms] STARTUP FAILED:")
    print(traceback.format_exc())

# INFINITE-FORMS-EOF
